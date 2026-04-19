# enhanced_table_analyzer.py
"""
增强的表格分析器 - 专门处理渠道政策表格数据
支持表格结构理解、自然语言查询、数值计算
"""

import pandas as pd
import re
import logging
import json
import os
from typing import Dict, List, Any, Optional, Union, Tuple
from dataclasses import dataclass, asdict
from docx import Document
from openpyxl import load_workbook

logger = logging.getLogger(__name__)

# 表格持久化文件路径
TABLES_CACHE_FILE = os.path.join(os.path.dirname(__file__), ".tables_cache.json")

# 导入数值表格增强器
try:
    from services.numeric_table_enhancer import get_numeric_table_enhancer
    USE_NUMERIC_ENHANCER = True
    logger.info("✅ 数值表格增强器已加载")
except ImportError:
    USE_NUMERIC_ENHANCER = False
    get_numeric_table_enhancer = None
    logger.warning("⚠️ 数值表格增强器不可用")


@dataclass
class TableInfo:
    """表格信息"""
    index: int
    rows: int
    cols: int
    headers: List[str]
    data: List[List[str]]
    table_type: str  # fee_schedule, product_list, comparison, etc.
    metadata: Dict[str, Any]


@dataclass
class QueryResult:
    """查询结果"""
    answer: str
    data: Optional[pd.DataFrame]
    confidence: float
    source_table: int
    explanation: str


class ChannelTableAnalyzer:
    """渠道表格分析器"""

    # 表格类型识别模式
    FEE_SCHEDULE_PATTERNS = [
        r'费用标准|费用编码|手续费|分成|激励',
        r'业务类型|费用|套餐|元',
        r'T\+\d+\s*月|T\d+|首充|递延|阶段'  # 新增：时间序列费用模式
    ]

    PRODUCT_LIST_PATTERNS = [
        r'产品名称|产品ID|优惠ID|产品编码',
        r'产品|优惠|ID|编码'
    ]

    COMPARISON_PATTERNS = [
        r'对比|差异|比较|变动',
        r'调整|变化|新增|取消'
    ]

    def __init__(self):
        """初始化表格分析器"""
        self.tables: List[TableInfo] = []
        # 启动时自动加载缓存的表格数据
        self._load_cached_tables()

        # 初始化数值表格增强器
        self.numeric_enhancer = None
        if USE_NUMERIC_ENHANCER:
            self.numeric_enhancer = get_numeric_table_enhancer()

        logger.info(f"渠道表格分析器初始化完成，已加载 {len(self.tables)} 个表格")

    def load_from_docx(self, docx_path: str) -> List[TableInfo]:
        """从 Word 文档加载表格"""
        try:
            doc = Document(docx_path)
            # 注意：不清空 self.tables，而是追加到现有表格列表
            # 这样可以支持从多个文档加载表格
            new_tables = []

            for i, table in enumerate(doc.tables):
                # 智能提取表头（处理双行表头）
                headers = []
                data_start_row = 1  # 默认从第1行开始是数据

                if len(table.rows) > 0:
                    # 提取第0行
                    row0_headers = [cell.text.strip() for cell in table.rows[0].cells]

                    # 检查是否有重复的表头（说明可能是双行表头）
                    header_counts = {}
                    for h in row0_headers:
                        header_counts[h] = header_counts.get(h, 0) + 1

                    has_duplicates = any(count > 1 for count in header_counts.values())

                    if has_duplicates and len(table.rows) > 1:
                        # 有重复表头，检查第1行是否有更具体的列名
                        row1_headers = [cell.text.strip() for cell in table.rows[1].cells]

                        # 合并表头：如果第0行和第1行都有内容，合并它们
                        merged_headers = []
                        for j in range(len(row0_headers)):
                            h0 = row0_headers[j] if j < len(row0_headers) else ""
                            h1 = row1_headers[j] if j < len(row1_headers) else ""

                            # 如果第1行更具体（包含更多信息），使用第1行
                            if h1 and len(h1) > len(h0) and h1 not in row0_headers:
                                merged_headers.append(h1)
                            # 如果第0行和第1行不同，合并它们
                            elif h0 and h1 and h0 != h1:
                                # 合并表头
                                merged_headers.append(f"{h0}_{h1}")
                            else:
                                # 使用非空的那个
                                merged_headers.append(h1 if h1 else h0)

                        headers = merged_headers
                        data_start_row = 2  # 数据从第2行开始
                    else:
                        # 没有重复表头，使用第0行作为表头
                        headers = row0_headers

                # 提取数据
                data = []
                for row in table.rows[data_start_row:]:  # 跳过表头行
                    row_data = [cell.text.strip() for cell in row.cells]
                    data.append(row_data)

                # 识别表格类型
                table_type = self._classify_table(headers, data)

                table_info = TableInfo(
                    index=len(self.tables) + len(new_tables),  # 使用全局索引
                    rows=len(table.rows),
                    cols=len(table.columns),
                    headers=headers,
                    data=data,
                    table_type=table_type,
                    metadata={
                        "source": docx_path,
                        "has_markings": self._has_boolean_markings(data)
                    }
                )

                new_tables.append(table_info)

            # 将新表格追加到现有表格列表
            self.tables.extend(new_tables)
            logger.info(f"从 {docx_path} 加载了 {len(new_tables)} 个表格，当前总共 {len(self.tables)} 个表格")
            return new_tables

        except Exception as e:
            logger.error(f"加载 Word 表格失败: {e}")
            return []

    def load_from_xlsx(self, xlsx_path: str, sheet_name: Optional[str] = None) -> List[TableInfo]:
        """从 Excel 文档加载表格"""
        try:
            wb = load_workbook(xlsx_path, data_only=True)

            tables = []
            sheet_names = [sheet_name] if sheet_name else wb.sheetnames

            for sheet_idx, sheet in enumerate(sheet_names):
                ws = wb[sheet]

                # 转换为 DataFrame
                data = []
                for row in ws.iter_rows(values_only=True):
                    data.append([str(cell) if cell is not None else "" for cell in row])

                if len(data) == 0:
                    continue

                # 假设第一行是表头
                headers = data[0]
                table_data = data[1:]

                # 识别表格类型
                table_type = self._classify_table(headers, table_data)

                table_info = TableInfo(
                    index=len(self.tables) + len(tables),  # 使用全局索引
                    rows=len(table_data) + 1,
                    cols=len(headers),
                    headers=headers,
                    data=table_data,
                    table_type=table_type,
                    metadata={
                        "source": xlsx_path,
                        "sheet": sheet
                    }
                )

                tables.append(table_info)

            self.tables.extend(tables)
            logger.info(f"从 {xlsx_path} 加载了 {len(tables)} 个表格，当前总共 {len(self.tables)} 个表格")
            return tables

        except Exception as e:
            logger.error(f"加载 Excel 表格失败: {e}")
            return []

    def _classify_table(self, headers: List[str], data: List[List[str]]) -> str:
        """分类表格类型"""
        headers_str = " ".join(headers)

        # 检查费用标准表
        for pattern in self.FEE_SCHEDULE_PATTERNS:
            if re.search(pattern, headers_str, re.IGNORECASE):
                return "fee_schedule"

        # 检查产品列表表
        for pattern in self.PRODUCT_LIST_PATTERNS:
            if re.search(pattern, headers_str, re.IGNORECASE):
                return "product_list"

        # 检查对比表
        for pattern in self.COMPARISON_PATTERNS:
            if re.search(pattern, headers_str, re.IGNORECASE):
                return "comparison"

        # 检查是否包含"√"或"X"标记
        if self._has_boolean_markings(data):
            return "matrix"

        return "general"

    def _has_boolean_markings(self, data: List[List[str]]) -> bool:
        """检查是否包含布尔标记（√/X）"""
        for row in data:
            for cell in row:
                if cell in ["√", "✓", "X", "×", "✔", "✗"]:
                    return True
        return False

    def to_dataframe(self, table_index: int) -> Optional[pd.DataFrame]:
        """将表格转换为 DataFrame"""
        if table_index >= len(self.tables):
            return None

        table = self.tables[table_index]
        df = pd.DataFrame(table.data, columns=table.headers)
        return df

    def query(self, natural_query: str, filters: Optional[Dict] = None) -> List[QueryResult]:
        """
        自然语言查询表格

        Args:
            natural_query: 自然语言查询
            filters: 过滤条件

        Returns:
            查询结果列表
        """
        results = []

        # 搜索相关表格
        relevant_tables = self._find_relevant_tables(natural_query)

        logger.info(f"📊 开始查询 {len(relevant_tables)} 个相关表格，查询: {natural_query[:50]}")

        for i, table_info in enumerate(relevant_tables):
            df = self.to_dataframe(table_info.index)

            if df is None:
                logger.warning(f"表格 {table_info.index} 转换为DataFrame失败")
                continue

            logger.info(f"📊 查询表格 {i} (index={table_info.index}), 形状: {df.shape}")

            # 使用智能查询方法
            result = self._smart_query(df, natural_query, table_info)
            if result:
                results.append(result)
                logger.info(f"✅ 表格 {table_info.index} 找到结果: {result.answer[:50]}")
            else:
                logger.info(f"❌ 表格 {table_info.index} 未找到匹配结果")

        logger.info(f"📊 表格查询完成，返回 {len(results)} 个结果")
        return results

    def _smart_query(self, df: pd.DataFrame, query: str, table_info: TableInfo) -> Optional[QueryResult]:
        """
        智能查询表格

        Args:
            df: 表格DataFrame
            query: 查询语句
            table_info: 表格信息

        Returns:
            查询结果
        """
        try:
            # 提取查询元素（保留数字和中文的组合）
            query_elements = self._extract_query_elements(query)
            logger.info(f"📊 查询元素: {query_elements[:10]}")

            # 检测表格类型（是否为转置表格，列名代表分类）
            is_transposed = self._is_transposed_table(df, query)

            if is_transposed:
                return self._query_transposed_table(df, query_elements, query, table_info)

            # 对每一行计算相关性分数
            scored_rows = []
            for idx, row in df.iterrows():
                row_text = " ".join([str(v) for v in row.values])
                score = self._calculate_relevance(row_text, query_elements, query)
                if score > 0:
                    scored_rows.append((idx, row_text, score))

            # 按分数排序
            scored_rows.sort(key=lambda x: x[2], reverse=True)

            if scored_rows:
                # 返回最高分的行，使用结构化格式确保字段对应清晰
                top_rows = scored_rows[:3]
                answer_parts = []

                # 智能检测相似字段组（如通用流量、定向流量）
                try:
                    similar_field_groups = self._detect_similar_fields(list(df.columns))
                except Exception as e:
                    logger.warning(f"相似字段检测失败: {e}")
                    similar_field_groups = []

                for idx, row_text, score in top_rows:
                    # 构建结构化的字段映射
                    row_data = df.iloc[idx]
                    structured_items = []

                    # 生成相似字段对比说明（通用方法）
                    field_comparison = self._generate_field_comparison(df, idx, similar_field_groups)

                    for header in df.columns:
                        header_str = str(header).strip()
                        value = str(row_data[header]).strip()

                        # 跳过空值
                        if not value or value == 'nan' or value == '/':
                            continue

                        # 使用清晰格式：字段名: 值
                        structured_items.append(f"{header_str}: {value}")

                    # 添加字段对比说明（优先显示）
                    if field_comparison:
                        answer_parts.append(field_comparison)

                    # 添加原始数据
                    answer_parts.extend(structured_items[:10])  # 限制显示前10个字段
                    answer_parts.append("")  # 空行分隔

                answer = "\n".join(answer_parts)

                # 数值增强：如果检测到数值密集型表格，添加结构化数据
                if self.numeric_enhancer:
                    answer = self.numeric_enhancer.enhance_table_query_result(answer, df)

                return QueryResult(
                    answer=answer,
                    data=df,
                    confidence=min(scored_rows[0][2], 1.0),
                    source_table=table_info.index,
                    explanation=f"找到 {len(scored_rows)} 个相关结果"
                )

            return None

        except Exception as e:
            logger.error(f"智能查询失败: {e}")
            return None

    def _is_transposed_table(self, df: pd.DataFrame, query: str) -> bool:
        """
        检测是否为转置表格（列名代表不同的分类）

        Args:
            df: 表格DataFrame
            query: 查询语句

        Returns:
            是否为转置表格
        """
        # 特征1: 检查列名中是否包含对比/分类关键词
        column_text = " ".join([str(col) for col in df.columns])

        # 转置表格的特征关键词（更广泛的检测）
        transposed_indicators = [
            "委托厅", "区域服务商", "非区域", "全业务",
            "服务商", "渠道", "校园", "微格", "自营",
            "129元", "套餐档位", "服补", "对比"
        ]

        count = sum(1 for indicator in transposed_indicators if indicator in column_text)
        if count >= 2:
            return True

        # 特征2: 检查表格是否有明显的对比结构
        # 如果第一列包含"倍数"、"比例"、"时间"等指标名称，且有多列数据
        if len(df.columns) >= 3:
            first_col_values = [str(val) for val in df.iloc[:, 0].dropna()]
            # 检查第一列是否包含指标名称
            metric_keywords = ["倍数", "比例", "时间", "发放", "档位", "门槛", "金额", "星级"]
            has_metrics = any(any(kw in val for kw in metric_keywords) for val in first_col_values)
            if has_metrics and len(df.columns) >= 3:
                return True

        return False

    def _query_transposed_table(self, df: pd.DataFrame, query_elements: List[str],
                                query: str, table_info: TableInfo) -> Optional[QueryResult]:
        """
        查询转置表格（列名代表不同分类）

        优化说明：
        - 包含所有列数据（不限于target_columns）
        - 增加行数限制到25行
        - 更清晰的列名对应关系

        Args:
            df: 表格DataFrame
            query_elements: 查询元素列表
            query: 原始查询
            table_info: 表格信息

        Returns:
            查询结果
        """
        try:
            # 获取列名并简化
            columns = [str(col) for col in df.columns]

            # 为列创建简化的别名（便于LLM理解）
            column_aliases = {}
            for i, col in enumerate(columns):
                if i == 0:
                    column_aliases[col] = "指标名称"
                elif "区域服务商" in col or ("区域" in col and "渠道" in col):
                    column_aliases[col] = "区域服务商"
                elif "委托厅" in col or "全业务" in col or "非区域" in col or "自营" in col or "Shopping" in col:
                    column_aliases[col] = "非区域服务商（委托厅）"
                else:
                    # 提取列名中的关键部分，保留更多关键信息
                    if len(col) > 25:
                        column_aliases[col] = col[:25]
                    else:
                        column_aliases[col] = col

            # 优化：获取所有有效数据列（不仅限于target_columns）
            all_data_columns = [col for col in columns[1:] if col and col.strip() and str(col) != 'nan']

            # 如果查询明确提到某些列，优先排列这些列
            priority_columns = []
            if any(kw in query for kw in ["区域服务商", "服务商"]):
                for col in all_data_columns:
                    if "区域" in col or ("服务商" in col and "非区域" not in col):
                        if col not in priority_columns:
                            priority_columns.append(col)

            if any(kw in query for kw in ["非区域", "委托厅", "全业务"]):
                for col in all_data_columns:
                    if any(kw in col for kw in ["委托", "全业务", "自营", "Shopping", "非区域"]):
                        if col not in priority_columns:
                            priority_columns.append(col)

            # 合并：优先列在前，其他列在后
            display_columns = priority_columns + [col for col in all_data_columns if col not in priority_columns]

            # 限制显示列数，避免过多影响可读性（最多12列）
            display_columns = display_columns[:12]

            # 构建比较结果
            result_parts = []

            # 添加表格结构说明 - 更清晰的格式
            structure_info = []
            structure_info.append("## 表格对比数据")
            structure_info.append("【数据说明】")
            structure_info.append("- 这是一个转置表格，第一列是指标名称，后续列是不同类别的数值")
            structure_info.append("- 下面列出每个指标在不同类别中的数值")
            structure_info.append(f"- 共{len(display_columns)}个类别，{len(df)}行数据")
            structure_info.append("")
            structure_info.append("【列名对应关系】")
            for i, col in enumerate(display_columns, 1):
                alias = column_aliases.get(col, col)
                structure_info.append(f"- 列{i}({alias}): {col}")
            structure_info.append("")
            structure_info.append("【对比数据详情】")
            result_parts.append("\n".join(structure_info))

            # 首先获取第一列作为行标签（通常是指标名称）
            first_col = df.columns[0] if len(df.columns) > 0 else None
            row_labels = []

            for idx, row in df.iterrows():
                if first_col:
                    label = str(row[first_col])
                    if label and label.strip() and label not in row_labels:
                        row_labels.append(label)

            # 优化：增加到25行，确保捕获所有相关数据
            for label in row_labels[:25]:
                matching_rows = df[df[first_col] == label]
                if len(matching_rows) > 0:
                    row = matching_rows.iloc[0]
                    row_values = [f"指标: {label}"]

                    # 遍历所有显示列
                    for col in display_columns:
                        val = str(row.get(col, ""))
                        if val and val != "nan" and val != "--":
                            alias = column_aliases.get(col, col)
                            row_values.append(f"列{display_columns.index(col)+1}({alias})={val}")

                    if len(row_values) > 1:  # 至少有一个值
                        result_parts.append(" | ".join(row_values))

            if result_parts:
                answer = "\n".join(result_parts)

                # 数值增强：如果检测到数值密集型表格，添加结构化数据
                if self.numeric_enhancer:
                    answer = self.numeric_enhancer.enhance_table_query_result(answer, df)

                return QueryResult(
                    answer=answer,
                    data=df,
                    confidence=0.9,
                    source_table=table_info.index,
                    explanation=f"从转置表格中提取 {min(len(row_labels), 25)} 行数据，涵盖 {len(display_columns)} 个类别"
                )

            return None

        except Exception as e:
            logger.error(f"查询转置表格失败: {e}")
            return None

    def _extract_query_elements(self, query: str) -> List[str]:
        """
        提取查询元素（保留有意义的组合）

        Args:
            query: 查询语句

        Returns:
            查询元素列表
        """
        elements = []

        # 模式1: 产品ID（如"prod.10086000050315"）- 优先提取
        # 支持 prod.xxx, product.xxx, ID 等格式
        product_id_patterns = re.findall(r'\b(?:prod|product|id)?\.?\s*[a-zA-Z]*\.?\s*\d{10,}', query, re.IGNORECASE)
        elements.extend(product_id_patterns)

        # 模式2: 数字+中文的组合（如"39元"、"5GB"、"6星"、"6星C"）
        # 修改为支持带字母后缀的中文组合
        patterns = re.findall(r'\d+[元GB级星级月笔C%天]+[A-Z]?', query)
        elements.extend(patterns)

        # 模式3: 时间段（如"T3-T12"、"T7-T12"）
        time_patterns = re.findall(r'T\d+(?:\s*[-~到]\s*T\d+)?', query, re.IGNORECASE)
        elements.extend(time_patterns)

        # 模式4: 英文字母+数字（如"T12"、"6星C"中的"C"）
        patterns = re.findall(r'[a-zA-Z]+\d+[\d.a-zA-Z]*|\d+[a-zA-Z]+', query)
        elements.extend(patterns)

        # 模式5: 纯数字（2位及以上）
        patterns = re.findall(r'\d{2,}', query)
        elements.extend(patterns)

        # 模式6: 中文词组（按长度排序，优先保留长词组）
        # 先提取长词组（4-10个字），再提取短词组（2-3个字）
        long_phrases = re.findall(r'[\u4e00-\u9fa5]{4,10}', query)
        short_phrases = re.findall(r'[\u4e00-\u9fa5]{2,3}', query)
        elements.extend(long_phrases)
        elements.extend(short_phrases)

        # 去重并保持顺序（长词组优先）
        seen = set()
        result = []
        for elem in elements:
            if elem not in seen:
                seen.add(elem)
                result.append(elem)

        return result

    def _calculate_relevance(self, row_text: str, query_elements: List[str], original_query: str) -> float:
        """
        计算行文本与查询的相关性分数

        Args:
            row_text: 行文本
            query_elements: 查询元素列表
            original_query: 原始查询

        Returns:
            相关性分数 (0-1)
        """
        score = 0.0
        row_text_lower = row_text.lower()

        # 精确匹配（完整元素匹配）给予高分
        for elem in query_elements:
            elem_lower = elem.lower()
            if elem_lower in row_text_lower:
                # 产品ID（如"prod.10086000050315"）- 最高权重
                if re.match(r'(?:prod|product)?\.?\s*\d{10,}', elem_lower):
                    score += 0.5
                # 时间段（如"T3-T12"）
                elif re.match(r'T\d+(?:\s*[-~到]\s*T\d+)?', elem, re.IGNORECASE):
                    score += 0.35
                # 数字+中文的组合（如"39元"、"6星"）- 高权重
                elif re.match(r'\d+[元GB级星级月笔C%天]+', elem):
                    score += 0.35
                # 英文字母+数字的组合（如"T12"）
                elif re.match(r'[a-zA-Z]+\d+', elem):
                    score += 0.3
                # 纯数字（2位以上）
                elif re.match(r'\d{2,}', elem):
                    score += 0.15
                # 长中文词组（4个字以上）
                elif len(elem) >= 4 and all('\u4e00' <= c <= '\u9fa5' for c in elem):
                    score += 0.25
                # 短中文词组
                else:
                    score += 0.1

        # 完整短语匹配（检查原始查询中的长短语，优先级更高）
        # 提取原始查询中的长短语（连续4个以上的中文）
        long_phrases = re.findall(r'[\u4e00-\u9fa5]{4,15}', original_query)
        for phrase in long_phrases:
            if phrase in row_text:
                # 长短语精确匹配给予很高权重
                score += 0.3

        # 短短语匹配（2-3个连续中文）
        short_phrases = re.findall(r'[\u4e00-\u9fa5]{2,3}', original_query)
        for phrase in short_phrases:
            if phrase in row_text:
                score += 0.05

        return min(score, 1.0)

    def _generate_field_comparison(self, df: pd.DataFrame, row_idx: int, similar_groups: List[Dict[str, Any]]) -> str:
        """
        生成相似字段对比说明（通用方法，帮助AI区分相似字段）

        Args:
            df: DataFrame
            row_idx: 行索引
            similar_groups: 相似字段组

        Returns:
            字段对比说明文本
        """
        if not similar_groups:
            return ""

        row_data = df.iloc[row_idx]
        comparison_items = []

        for group in similar_groups:
            members = group['members']
            if len(members) < 2:
                continue

            # 提取这些字段在当前行的值
            values = []
            for member in members:
                if member in df.columns:
                    value = str(row_data[member]).strip()
                    # 清理字段名（去除单位、括号等）
                    clean_name = re.sub(r'[（(].*?[）)]', '', member).strip()
                    values.append(f"{clean_name}={value}")

            # 如果这些字段的值不同，添加到对比说明中
            if len(values) >= 2:
                # 检查值是否相同
                unique_values = list(set([v.split('=')[1] for v in values]))
                if len(unique_values) > 1:
                    comparison_items.append(", ".join(values))

        if comparison_items:
            return f"[字段对比] {', '.join(comparison_items)}"
        else:
            return ""

    def _detect_similar_fields(self, headers: List[str]) -> List[Dict[str, Any]]:
        """
        检测相似的字段名（通用算法，不硬编码）

        策略：
        1. 提取字段名的核心词汇（去除括号、单位等）
        2. 计算字段之间的相似度
        3. 将相似字段分组

        Args:
            headers: 表头列表

        Returns:
            相似字段组列表，每组包含组名和成员
        """
        if not headers:
            return []

        # 确保 headers 都是字符串
        headers_str = [str(h).strip() for h in headers]

        # 1. 清理字段名，提取核心词汇
        cleaned_fields = {}
        for header in headers_str:
            header_str = str(header).strip()
            # 去除括号内容（单位等）
            core_name = re.sub(r'[（(].*?[）)]', '', header_str)
            # 去除常见后缀
            core_name = re.sub(r'(?:流量|资费|费用|服务)$', '', core_name)
            cleaned_fields[header_str] = core_name

        # 2. 检测相似字段组
        field_groups = []
        processed = set()

        for header1, core1 in cleaned_fields.items():
            if header1 in processed:
                continue

            group_members = [header1]

            # 查找与当前字段相似的其他字段
            for header2, core2 in cleaned_fields.items():
                if header2 == header1 or header2 in processed:
                    continue

                # 计算相似度
                similarity = self._calculate_field_similarity(core1, core2)

                # 如果相似度超过阈值，认为是同一组
                if similarity >= 0.6:
                    group_members.append(header2)
                    processed.add(header2)

            # 如果组内有多个成员，记录下来
            if len(group_members) > 1:
                # 提取共同前缀作为组名
                group_name = self._extract_group_name(group_members)
                field_groups.append({
                    'name': group_name,
                    'members': group_members
                })

            processed.add(header1)

        return field_groups

    def _calculate_field_similarity(self, field1: str, field2: str) -> float:
        """
        计算两个字段的相似度

        Args:
            field1: 字段1名称
            field2: 字段2名称

        Returns:
            相似度 (0-1)
        """
        # 完全相同
        if field1 == field2:
            return 1.0

        # 包含关系
        if field1 in field2 or field2 in field1:
            return 0.8

        # 计算编辑距离
        distance = self._levenshtein_distance(field1, field2)
        max_len = max(len(field1), len(field2))
        similarity = 1 - (distance / max_len) if max_len > 0 else 0

        return similarity

    def _levenshtein_distance(self, s1: str, s2: str) -> int:
        """计算编辑距离"""
        if len(s1) < len(s2):
            return self._levenshtein_distance(s2, s1)

        if len(s2) == 0:
            return len(s1)

        previous_row = range(len(s2) + 1)
        for i, c1 in enumerate(s1):
            current_row = [i + 1]
            for j, c2 in enumerate(s2):
                insertions = previous_row[j + 1] + 1
                deletions = current_row[j] + 1
                substitutions = previous_row[j] + (c1 != c2)
                current_row.append(min(insertions, deletions, substitutions))
            previous_row = current_row

        return previous_row[-1]

    def _extract_group_name(self, members: List[str]) -> str:
        """
        从相似字段成员中提取组名

        策略：提取共同的前缀或核心词汇

        Args:
            members: 字段成员列表

        Returns:
            组名
        """
        if not members:
            return ""

        # 清理成员名称
        cleaned = []
        for member in members:
            # 去除括号内容
            clean = re.sub(r'[（(].*?[）)]', '', member)
            # 去除空格
            clean = clean.strip()
            cleaned.append(clean)

        # 找出共同前缀
        if len(cleaned) == 1:
            return cleaned[0]

        # 使用第一个成员作为基准
        first = cleaned[0]
        common_prefix = first

        for member in cleaned[1:]:
            # 逐步缩短共同前缀
            while not member.startswith(common_prefix) and len(common_prefix) > 0:
                common_prefix = common_prefix[:-1]

        # 如果共同前缀太短（少于2个字符），使用第一个成员
        if len(common_prefix) < 2:
            return first

        return common_prefix

    def _get_field_group_info(self, field_name: str, groups: List[Dict[str, Any]]) -> Optional[Tuple[str, List[str]]]:
        """
        获取字段所属的组信息

        Args:
            field_name: 字段名
            groups: 相似字段组列表

        Returns:
            (组名, 组成员) 或 None
        """
        # 确保 field_name 是字符串
        field_name_str = str(field_name).strip()

        for group in groups:
            if field_name_str in group['members']:
                return (group['name'], group['members'])
        return None

    def _parse_query_intent(self, query: str) -> str:
        """解析查询意图"""
        if any(kw in query for kw in ["多少钱", "费用", "标准", "多少元"]):
            return "fee_query"
        elif any(kw in query for kw in ["适用", "可以用", "哪些"]):
            return "applicability"
        elif any(kw in query for kw in ["产品", "套餐", "ID"]):
            return "product_lookup"
        elif any(kw in query for kw in ["对比", "差异", "不同"]):
            return "comparison"
        else:
            return "general"

    def _find_relevant_tables(self, query: str) -> List[TableInfo]:
        """查找相关表格（通用化，无硬编码业务术语）"""
        relevant = []

        # 通用关键词提取（不绑定具体业务）
        keywords = []

        # 1. 提取长中文词组（优先级高，用于精确匹配）
        long_phrases = re.findall(r'[\u4e00-\u9fa5]{4,10}', query)
        keywords.extend(long_phrases)

        # 2. 提取短中文词组
        short_phrases = re.findall(r'[\u4e00-\u9fa5]{2,3}', query)
        keywords.extend(short_phrases)

        # 3. 英文单词（3个字母以上）
        english_words = re.findall(r'[a-zA-Z]{3,}', query)
        keywords.extend(english_words)

        # 4. 数字+特殊符号组合（产品ID、时间段等）
        special_patterns = re.findall(r'[a-zA-Z]+\d+[\d.a-zA-Z]*|\d+[a-zA-Z%]+|T\d+|prod\.\d+|product\.\d+', query, re.IGNORECASE)
        keywords.extend(special_patterns)

        # 5. 数字+中文模式
        number_chinese = re.findall(r'\d+[元GB级星级月笔C%天]+[A-Z]?', query)
        keywords.extend(number_chinese)

        # 去重
        keywords = list(set(keywords))

        logger.info(f"📊 表格查询关键词: {keywords[:15]}")  # 记录前15个关键词用于调试

        for table in self.tables:
            # 计算相关性分数
            score = 0
            # 扩展表格文本匹配范围（包含更多行数据和表头）
            table_text = " ".join(table.headers) + " " + " ".join([" ".join(row) for row in table.data[:20]])

            for keyword in keywords:
                keyword_lower = keyword.lower()
                table_text_lower = table_text.lower()

                if keyword_lower in table_text_lower:
                    # 长关键词给予更高分数
                    if len(keyword) >= 4:
                        score += 2
                    else:
                        score += 1

            # 额外加分：检查表头是否包含关键词
            header_text = " ".join(table.headers).lower()
            for keyword in keywords:
                if keyword.lower() in header_text:
                    score += 3  # 表头匹配权重更高

            if score > 0:
                relevant.append((table, score))

        # 按相关性排序
        relevant.sort(key=lambda x: x[1], reverse=True)

        result = [t[0] for t in relevant[:5]]  # 返回前5个相关表格
        logger.info(f"📊 找到 {len(result)} 个相关表格，得分: {[t[1] for t in relevant[:5]]}")
        return result

    def _execute_query(self, df: pd.DataFrame, query: str,
                      intent: str, table_info: TableInfo) -> Optional[QueryResult]:
        """执行查询"""

        try:
            if intent == "fee_query":
                return self._query_fee(df, query, table_info)
            elif intent == "applicability":
                return self._query_applicability(df, query, table_info)
            elif intent == "product_lookup":
                return self._query_product(df, query, table_info)
            else:
                return self._query_general(df, query, table_info)

        except Exception as e:
            logger.error(f"执行查询失败: {e}")
            return None

    def _query_fee(self, df: pd.DataFrame, query: str, table_info: TableInfo) -> Optional[QueryResult]:
        """查询费用"""

        # 提取金额信息
        amounts = re.findall(r'(\d+(?:\.\d+)?)\s*(?:元|块)', query)
        fee_types = re.findall(r'(实名|充值|分成|手续费|激励)', query)

        # 简单实现：在表格中搜索匹配的行
        answer_parts = []

        for _, row in df.iterrows():
            row_text = " ".join([str(v) for v in row.values])

            # 检查是否匹配查询条件
            match = True
            if amounts:
                if not any(amt in row_text for amt in amounts):
                    match = False

            if fee_types:
                if not any(ft in row_text for ft in fee_types):
                    match = False

            if match:
                answer_parts.append(f"第{len(answer_parts)+1}行: {row_text[:100]}")

        if answer_parts:
            answer = "\n".join(answer_parts[:3])  # 返回前3个匹配
            return QueryResult(
                answer=answer,
                data=df,
                confidence=0.8,
                source_table=table_info.index,
                explanation=f"从{table_info.table_type}表中找到{len(answer_parts)}个匹配结果"
            )

        return None

    def _query_applicability(self, df: pd.DataFrame, query: str, table_info: TableInfo) -> Optional[QueryResult]:
        """查询适用性"""

        # 查找包含"√"的单元格
        answer = ""
        count = 0

        for col_idx, col_name in enumerate(df.columns):
            # 检查该列的条件（通常在第一列）
            for idx, row in df.iterrows():
                cell_value = str(row.iloc[col_idx])

                if cell_value in ["√", "✓", "✔"]:
                    condition = str(row.iloc[0])  # 假设第一列是条件
                    applicable = col_name
                    answer += f"- 当{condition}时，{applicable}适用\n"
                    count += 1

        if answer:
            return QueryResult(
                answer=answer.strip(),
                data=df,
                confidence=0.85,
                source_table=table_info.index,
                explanation=f"找到{count}个适用规则"
            )

        return None

    def _query_product(self, df: pd.DataFrame, query: str, table_info: TableInfo) -> Optional[QueryResult]:
        """查询产品"""

        # 提取产品名称或ID
        product_names = re.findall(r'[\u4e00-\u9fa5]{2,}', query)
        product_ids = re.findall(r'[A-Z]{2}\d{6,}', query, re.IGNORECASE)

        matches = []

        for _, row in df.iterrows():
            row_text = " ".join([str(v) for v in row.values])

            # 匹配产品名称
            for name in product_names:
                if name in row_text:
                    matches.append(f"找到产品: {row_text[:80]}")
                    break

            # 匹配产品ID
            for pid in product_ids:
                if pid.upper() in row_text.upper():
                    matches.append(f"找到ID {pid}: {row_text[:80]}")
                    break

        if matches:
            return QueryResult(
                answer="\n".join(matches[:3]),
                data=df,
                confidence=0.9,
                source_table=table_info.index,
                explanation=f"找到{len(matches)}个产品匹配"
            )

        return None

    def _query_general(self, df: pd.DataFrame, query: str, table_info: TableInfo) -> Optional[QueryResult]:
        """通用查询"""

        # 简单的关键词匹配
        keywords = re.findall(r'[\u4e00-\u9fa5]{2,}', query)

        matches = []
        for _, row in df.iterrows():
            row_text = " ".join([str(v) for v in row.values])
            match_count = sum(1 for kw in keywords if kw in row_text)

            if match_count >= 2:  # 至少匹配2个关键词
                matches.append((row_text[:100], match_count))

        if matches:
            matches.sort(key=lambda x: x[1], reverse=True)
            answer = "\n".join([m[0] for m in matches[:3]])

            return QueryResult(
                answer=answer,
                data=df,
                confidence=0.7,
                source_table=table_info.index,
                explanation=f"根据关键词匹配找到{len(matches)}个结果"
            )

        return None

    def get_table_summary(self) -> List[Dict[str, Any]]:
        """获取所有表格的摘要信息"""
        summary = []

        for table in self.tables:
            summary.append({
                "index": table.index,
                "type": table.table_type,
                "rows": table.rows,
                "cols": table.cols,
                "headers": table.headers[:5],  # 只显示前5个列名
                "source": table.metadata.get("source", "")
            })

        return summary

    def _generate_table_keywords(self, table: 'TableData') -> List[str]:
        """
        为表格生成关键词，用于提高检索匹配度

        Args:
            table: 表格数据对象

        Returns:
            关键词列表
        """
        keywords = []

        try:
            # 转换为DataFrame便于处理
            import pandas as pd
            df = pd.DataFrame(table.data)

            # 1. 检查是否是套餐对比表
            if '套餐名称' in table.headers or '产品ID' in table.headers:
                for _, row in df.iterrows():
                    if len(row) > 0:
                        plan_name = str(row.iloc[0])
                        # 提取套餐名称和价格
                        if '潮玩青春卡' in plan_name:
                            keywords.append(f"{plan_name}套餐详情")

                            # 提取流量信息
                            for idx, val in enumerate(row):
                                val_str = str(val)
                                header = table.headers[idx] if idx < len(table.headers) else ""
                                if 'G' in val_str and '通用' in header:
                                    keywords.append(f"{plan_name}包含{val_str}通用流量")
                                elif 'G' in val_str and '定向' in header:
                                    keywords.append(f"{plan_name}包含{val_str}定向流量")

            # 2. 检查是否是费用标准表
            elif '费用标准' in table.headers or '费用编码' in table.headers:
                for _, row in df.iterrows():
                    if len(row) > 0:
                        fee_type = str(row.iloc[0])
                        keywords.append(f"{fee_type}费用标准")

                        # 提取费用金额
                        for val in row:
                            val_str = str(val)
                            if '元' in val_str and any(c.isdigit() for c in val_str):
                                keywords.append(f"{fee_type}{val_str}")

            # 3. 检查是否是激励规则表
            elif any(keyword in ' '.join(table.headers) for keyword in ['激励', '分成', '奖励']):
                keywords.append("酬金政策")
                keywords.append("分成激励规则")
                keywords.append("渠道激励标准")

            # 4. 通用关键词：从表头生成
            for header in table.headers[:5]:  # 只取前5个表头
                if header and len(header) > 2:
                    keywords.append(header)

        except Exception as e:
            logger.warning(f"生成表格关键词时出错: {e}")

        return list(set(keywords))  # 去重

    def save_tables_to_cache(self) -> bool:
        """
        保存表格数据到缓存文件

        Returns:
            是否保存成功
        """
        import tempfile
        import shutil

        try:
            logger.info(f"💾 开始保存表格缓存，表格数量: {len(self.tables)}")

            # 将表格数据转换为可序列化的格式
            tables_data = []
            failed_tables = []

            for idx, table in enumerate(self.tables):
                try:
                    # 生成表格关键词（用于检索优化）
                    keywords = self._generate_table_keywords(table)

                    # 尝试序列化单个表格，检查是否有不可序列化的数据
                    table_dict = {
                        "index": table.index,
                        "rows": table.rows,
                        "cols": table.cols,
                        "headers": table.headers,
                        "data": table.data,
                        "table_type": table.table_type,
                        "metadata": table.metadata,
                        "keywords": keywords  # 添加关键词字段
                    }
                    # 测试序列化
                    json.dumps(table_dict, ensure_ascii=False)
                    tables_data.append(table_dict)

                except (TypeError, ValueError) as e:
                    logger.warning(f"⚠️ 表格 {idx} 数据无法序列化，跳过: {e}")
                    failed_tables.append(idx)
                    continue

            if failed_tables:
                logger.warning(f"⚠️ 跳过了 {len(failed_tables)} 个无法序列化的表格")

            # 准备缓存数据
            cache_data = {
                "version": "1.0",
                "tables_count": len(tables_data),
                "tables": tables_data
            }

            # 先测试完整序列化
            try:
                json_str = json.dumps(cache_data, ensure_ascii=False, indent=2)
                logger.info(f"✅ 序列化成功，数据大小: {len(json_str)} 字符")
            except Exception as e:
                logger.error(f"❌ 序列化失败: {e}")
                return False

            # 使用临时文件进行原子性写入
            temp_file = TABLES_CACHE_FILE + '.tmp'
            try:
                with open(temp_file, 'w', encoding='utf-8') as f:
                    f.write(json_str)

                # 原子性替换文件
                shutil.move(temp_file, TABLES_CACHE_FILE)

                logger.info(f"✅ 表格数据已保存到缓存: {len(tables_data)} 个表格")
                return True

            except Exception as e:
                logger.error(f"❌ 写入文件失败: {e}")
                # 清理临时文件
                if os.path.exists(temp_file):
                    try:
                        os.remove(temp_file)
                    except:
                        pass
                return False

        except Exception as e:
            logger.error(f"❌ 保存表格缓存失败: {e}")
            import traceback
            logger.error(f"详细错误堆栈:\n{traceback.format_exc()}")
            return False

    def _load_cached_tables(self) -> bool:
        """
        从缓存文件加载表格数据

        Returns:
            是否加载成功
        """
        try:
            if not os.path.exists(TABLES_CACHE_FILE):
                logger.info("📭 表格缓存文件不存在，将从零开始")
                return False

            with open(TABLES_CACHE_FILE, 'r', encoding='utf-8') as f:
                cache_data = json.load(f)

            tables_data = cache_data.get("tables", [])
            loaded_count = 0

            for table_dict in tables_data:
                try:
                    table_info = TableInfo(
                        index=table_dict["index"],
                        rows=table_dict["rows"],
                        cols=table_dict["cols"],
                        headers=table_dict["headers"],
                        data=table_dict["data"],
                        table_type=table_dict["table_type"],
                        metadata=table_dict["metadata"]
                    )
                    self.tables.append(table_info)
                    loaded_count += 1
                except Exception as e:
                    logger.warning(f"⚠️ 加载单个表格失败: {e}")

            logger.info(f"✅ 从缓存加载了 {loaded_count} 个表格")
            return loaded_count > 0

        except Exception as e:
            logger.error(f"❌ 加载表格缓存失败: {e}")
            return False


# 单例模式
_table_analyzer = None

def get_table_analyzer() -> ChannelTableAnalyzer:
    """获取表格分析器实例"""
    global _table_analyzer
    if _table_analyzer is None:
        _table_analyzer = ChannelTableAnalyzer()
    return _table_analyzer
