"""
重新处理和索引文档，更新向量数据库
"""

import sys
import os
sys.path.insert(0, os.path.dirname(__file__))

def reindex_documents():
    """重新处理文档，更新向量数据库"""
    # 设置离线模式，使用本地缓存的模型
    import os
    os.environ['HF_HUB_OFFLINE'] = '1'
    os.environ['TRANSFORMERS_OFFLINE'] = '1'

    print("=" * 80)
    print("重新处理和索引文档")
    print("=" * 80)

    # 导入必要的模块
    try:
        from qdrant_client import QdrantClient
        from config import VectorConfig
        from embedding_service import EmbeddingService
        from docx import Document
        import logging

        logging.basicConfig(level=logging.INFO)

        print("\n[步骤1] 连接Qdrant向量数据库")
        print("-" * 80)

        config = VectorConfig.from_env()
        client = QdrantClient(
            host=config.qdrant_host,
            port=config.qdrant_port,
            timeout=10
        )

        # 检查collection
        try:
            collections = client.get_collections().collections
            collection_names = [c.name for c in collections]
            print(f"现有collections: {collection_names}")

            # 删除旧的documents collection
            if "documents" in collection_names:
                print("正在删除旧的documents collection...")
                client.delete_collection("documents")
                print("[OK] 已删除旧collection")
        except Exception as e:
            print(f"[INFO] {e}")

        print("\n[步骤2] 处理Word文档")
        print("-" * 80)

        doc_path = 'c:/Users/Administrator/Desktop/12月渠道政策/12月放号独立部分（政策-操作-ID）.docx'

        if not os.path.exists(doc_path):
            print(f"[ERROR] 文档不存在: {doc_path}")
            return

        # 读取文档
        doc = Document(doc_path)
        print(f"文档共有 {len(doc.tables)} 个表格")

        # 提取所有文本片段（包括表格）
        text_chunks = []

        for table_idx, table in enumerate(doc.tables):
            # 提取表格标题（如果有）
            table_context = f"表格{table_idx + 1}: "

            # 提取表头
            if len(table.rows) > 1:
                headers = [cell.text.strip() for cell in table.rows[1].cells]
                table_context += " | ".join(headers[:10])

            text_chunks.append(table_context)

            # 提取数据行
            for row_idx, row in enumerate(table.rows[2:], start=2):  # 跳过表头行
                row_text = " | ".join([cell.text.strip()[:30] for cell in row.cells[:10]])
                if row_text.strip():
                    text_chunks.append(row_text)

        print(f"提取了 {len(text_chunks)} 个文本片段")

        print("\n[步骤3] 生成向量并索引")
        print("-" * 80)

        # 初始化嵌入服务
        embed_service = EmbeddingService()
        if not embed_service.load_model():
            print("[ERROR] 无法加载嵌入模型")
            return

        # 创建collection
        from qdrant_client.http import models

        client.create_collection(
            collection_name="documents",
            vectors_config=models.VectorParams(
                size=512,  # 嵌入向量维度
                distance=models.Distance.COSINE
            )
        )
        print("[OK] 已创建新collection")

        # 批量插入数据
        batch_size = 10
        total_inserted = 0

        for i in range(0, len(text_chunks), batch_size):
            batch = text_chunks[i:i+batch_size]

            # 生成向量
            vectors = [embed_service.embed_text(text) for text in batch]

            # 准备payload
            payloads = [{"text": text, "source": doc_path} for text in batch]

            # 插入数据
            client.upsert(
                collection_name="documents",
                points=models.Batch(
                    ids=[total_inserted + j for j in range(len(batch))],
                    vectors=vectors,
                    payloads=payloads
                )
            )

            total_inserted += len(batch)
            print(f"已插入 {total_inserted}/{len(text_chunks)} 个片段")

        print(f"\n[OK] 成功索引 {total_inserted} 个文本片段")

        # 验证索引
        print("\n[步骤4] 验证索引")
        print("-" * 80)

        query = "59元潮玩青春卡套餐包含哪些内容？"
        query_vector = embed_service.embed_text(query)

        search_results = client.search(
            collection_name="documents",
            query_vector=query_vector,
            limit=3,
            score_threshold=0.3
        )

        print(f"查询测试: {query}")
        print(f"找到 {len(search_results)} 个结果")

        for i, hit in enumerate(search_results[:2], 1):
            print(f"\n结果 {i} (得分: {hit.score:.2f}):")
            if hasattr(hit, 'payload'):
                print(f"  {hit.payload.get('text', '')[:100]}...")

    except Exception as e:
        print(f"\n[ERROR] 重新索引失败: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    try:
        reindex_documents()
        print("\n" + "=" * 80)
        print("[完成] 重新索引完成")
        print("=" * 80)
    except Exception as e:
        print(f"\n[错误] {e}")
