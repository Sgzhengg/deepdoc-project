# main.py - PaddleOCR版本
import io
import json
import logging
import os
import sys
import tempfile
import atexit
import uuid
import datetime
from typing import Dict, Any, Optional, Tuple, List
from pathlib import Path
from contextlib import asynccontextmanager

# 加载 .env 文件（必须在最开始）
from dotenv import load_dotenv
load_dotenv()

# 导入服务实例模块
import service_instance

# 配置日志（必须在最开始）
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# 配置镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'
# 禁用模型源检查以加快启动
os.environ['DISABLE_MODEL_SOURCE_CHECK'] = 'True'

import uvicorn
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, BackgroundTasks, Query, Request
from fastapi.responses import JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.middleware.gzip import GZipMiddleware
from fastapi.encoders import jsonable_encoder
import json

# Office文档处理
try:
    import pandas as pd
    from docx import Document as DocxDocument
    OFFICE_SUPPORT = True
except ImportError:
    OFFICE_SUPPORT = False
    logger.warning("pandas 或 python-docx 未安装")

# PDF处理
try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False

# PaddleOCR导入（PPStructure 在某些版本可能不可用）
try:
    from paddleocr import PaddleOCR
    try:
        from paddleocr import PPStructure
        PPSTRUCTURE_AVAILABLE = True
    except ImportError:
        PPSTRUCTURE_AVAILABLE = False
        logger.warning("PPStructure 不可用，仅使用基本 OCR 功能")
    PADDLEOCR_AVAILABLE = True
except ImportError:
    PADDLEOCR_AVAILABLE = False
    PPSTRUCTURE_AVAILABLE = False
    logger.warning("PaddleOCR未安装，OCR功能将不可用")

# ========== 创建FastAPI应用 ==========

# 全局服务实例
_chat_service = None  # 旧版单模型服务（Ollama，向后兼容）
_dual_model_service = None  # 双模型服务
_single_model_service = None  # 新版单模型服务（本地AWQ模型）
_embedding_service = None
_vector_storage = None

@asynccontextmanager
async def lifespan(app: FastAPI):
    """应用生命周期管理"""
    # 启动时初始化服务
    logger.info("🚀 DeepDoc API启动中...")

    # 从环境变量读取配置
    ollama_base_url = os.getenv("OLLAMA_BASE_URL", "http://localhost:11434")
    model_name = os.getenv("LLM_MODEL", "deepseek-r1:32b")
    use_deepseek_api = os.getenv("USE_DEEPSEEK_API", "false").lower() == "true"

    # 如果使用 DeepSeek API，跳过本地 ChatService 初始化
    if not use_deepseek_api:
        # 使用单模型服务（只使用deepseek-r1:32B进行推理）
        try:
            from services.chat_service import get_chat_service
            chat_service = get_chat_service()

            chat_service.initialize(
                ollama_base_url=ollama_base_url,
                model_name=model_name
            )

            global _chat_service
            _chat_service = chat_service
            sys.modules['_chat_service_instance'] = chat_service
            service_instance.set_chat_service(chat_service)

            logger.info("✅ ChatService 初始化完成")
            logger.info(f"📊 推理模型: {model_name}")

        except Exception as e:
            logger.error(f"❌ ChatService 初始化失败: {e}")
            import traceback
            logger.error(traceback.format_exc())
            logger.warning("⚠️  聊天功能可能不可用，但其他API仍可正常工作")
    else:
        logger.info("✅ 使用 DeepSeek API 模式，跳过本地 ChatService 初始化")
        logger.info("📊 推理模型: DeepSeek API (云端)")

    # ========== 预加载嵌入服务 ==========
    global _embedding_service
    try:
        from embedding_service import EmbeddingService

        logger.info("📦 预加载嵌入服务...")
        _embedding_service = EmbeddingService()
        _embedding_service.load_model()

        # 将嵌入服务存储到全局模块中，方便其他模块访问
        sys.modules['_embedding_service_instance'] = _embedding_service

        logger.info("✅ 嵌入服务预加载完成")
        logger.info(f"   模型: {_embedding_service.model_info.get('name')}")
        logger.info(f"   维度: {_embedding_service.model_info.get('dimension')}")
        logger.info(f"   设备: {_embedding_service.model_info.get('device')}")

    except Exception as e:
        logger.error(f"❌ 嵌入服务预加载失败: {e}")
        import traceback
        logger.error(traceback.format_exc())
        logger.warning("⚠️  向量检索功能可能不可用")
        _embedding_service = None

    # ========== 初始化向量存储服务 ==========
    global _vector_storage
    try:
        from vector_storage import VectorStorage

        logger.info("📦 初始化向量存储服务...")
        _vector_storage = VectorStorage()
        _vector_storage.connect()

        # 将向量存储服务存储到全局模块中，方便其他模块访问
        sys.modules['_vector_storage_instance'] = _vector_storage

        logger.info("✅ 向量存储服务初始化完成")

    except Exception as e:
        logger.error(f"❌ 向量存储服务初始化失败: {e}")
        import traceback
        logger.error(traceback.format_exc())
        logger.warning("⚠️  文档上传和检索功能可能不可用")
        _vector_storage = None

    yield  # 应用运行

    # ========== 关闭时清理资源 ==========
    logger.info("👋 DeepDoc API关闭中...")

    # 清理全局服务实例
    if _chat_service:
        try:
            logger.info("🧹 清理 ChatService...")
            # 如果有 cleanup 方法，调用它
            if hasattr(_chat_service, 'cleanup'):
                _chat_service.cleanup()
            _chat_service = None
            logger.info("✅ ChatService 已清理")
        except Exception as e:
            logger.error(f"❌ 清理 ChatService 失败: {e}")

    if _embedding_service:
        try:
            logger.info("🧹 清理嵌入服务...")
            _embedding_service = None
            logger.info("✅ 嵌入服务已清理")
        except Exception as e:
            logger.error(f"❌ 清理嵌入服务失败: {e}")

    if _vector_storage:
        try:
            logger.info("🧹 清理向量存储...")
            # 如果有 close 方法，调用它
            if hasattr(_vector_storage, 'close'):
                _vector_storage.close()
            _vector_storage = None
            logger.info("✅ 向量存储已清理")
        except Exception as e:
            logger.error(f"❌ 清理向量存储失败: {e}")

    # 清理系统模块中的引用
    for module_name in ['_chat_service_instance', '_embedding_service_instance', '_vector_storage_instance']:
        if module_name in sys.modules:
            del sys.modules[module_name]

    logger.info("✅ DeepDoc API 已安全关闭")

app = FastAPI(
    title="DeepDoc API - LangGraph多Agent版本",
    description="基于LangGraph的智能文档管理系统，支持Agentic RAG",
    version="4.0.0",
    lifespan=lifespan
)

# 自定义JSON响应类，确保UTF-8编码
class UTF8JSONResponse(JSONResponse):
    def render(self, content: Any) -> bytes:
        return json.dumps(
            jsonable_encoder(content),
            ensure_ascii=False,
            allow_nan=True,
            indent=None,
            separators=(",", ":"),
        ).encode("utf-8")

# CORS中间件 - 从环境变量读取允许的域名
# 开发环境: * (允许所有来源)
# 生产环境: 应设置具体域名，如 "https://your-mobile-app.com,https://your-pc-app.com"
allowed_origins_str = os.getenv("ALLOWED_ORIGINS", "*")
allowed_origins = allowed_origins_str.split(",") if allowed_origins_str != "*" else ["*"]

# 当使用credentials时，不能使用通配符，需要明确指定允许的源
if "*" in allowed_origins:
    # 开发环境：允许所有localhost端口以及常见测试域名
    allowed_origins = [
        "http://localhost:3000",
        "http://localhost:8080",
        "http://127.0.0.1:3000",
        "http://127.0.0.1:8080",
        "http://localhost:5173",
        "http://127.0.0.1:5173",
        # 花生壳公网测试域名（前端和后端）
        "http://1203862ikgl90.vicp.fun",
        "http://1203862ikgl90.vicp.fun:8080",
    ]

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
    expose_headers=["*"],
)
app.add_middleware(GZipMiddleware, minimum_size=1000)


# ========== API密钥认证中间件（可选）==========
API_KEY = os.getenv("API_KEY", None)  # 从环境变量读取API密钥

@app.middleware("http")
async def api_key_middleware(request, call_next):
    """API密钥认证中间件（如果设置了API_KEY环境变量）"""
    # 如果未设置API密钥，跳过认证
    if not API_KEY:
        return await call_next(request)

    # 健康检查和公开接口不需要认证
    if request.url.path in ["/health", "/docs", "/openapi.json"]:
        return await call_next(request)

    # 检查请求头中的API密钥
    client_api_key = request.headers.get("X-API-Key")
    if client_api_key != API_KEY:
        return JSONResponse(
            status_code=403,
            content={"detail": "Invalid or missing API Key"}
        )

    return await call_next(request)

# ========== 导入向量检索模块 ==========
try:
    from embedding_service import EmbeddingService
    from vector_storage import VectorStorage
    from hybrid_search import HybridRetriever

    VECTOR_FEATURES_ENABLED = True
    HYBRID_SEARCH_ENABLED = True
    logger.info("✅ 向量检索和混合检索功能已启用")

except ImportError as e:
    logger.warning(f"⚠️  无法导入向量检索模块: {e}")
    VECTOR_FEATURES_ENABLED = False
    HYBRID_SEARCH_ENABLED = False

    class EmbeddingService:
        def __init__(self): pass
        def load_model(self): return False
        def get_model_info(self): return {"loaded": False}

    class VectorStorage:
        def __init__(self): pass
        def connect(self): return False

    class HybridRetriever:
        def __init__(self, *args, **kwargs): pass

# ========== 导入其他模块 ==========
try:
    from kb_management import KnowledgeBaseManager
    KB_MANAGEMENT_ENABLED = True
    logger.info("✅ 知识库管理功能已启用")
except ImportError:
    KB_MANAGEMENT_ENABLED = False

# ========== 全局服务实例 ==========
_paddle_ocr = None
_pp_structure = None
_embedding_service = None
_vector_storage = None
_hybrid_retriever = None
_kb_manager = None

# ========== 辅助函数 ==========
def safe_extract(obj, attr_name, default=None):
    """安全提取属性"""
    try:
        return getattr(obj, attr_name, default)
    except Exception:
        return default

# ========== PaddleOCR服务函数 ==========
def get_paddle_ocr(use_structure=False):
    """获取PaddleOCR实例"""
    global _paddle_ocr, _pp_structure

    if not PADDLEOCR_AVAILABLE:
        raise HTTPException(status_code=500, detail="PaddleOCR未安装")

    if use_structure:
        # 检查PPStructure是否可用
        if not PPSTRUCTURE_AVAILABLE:
            logger.warning("PPStructure 不可用，降级到普通 OCR 模式")
            return get_paddle_ocr(use_structure=False)
        if _pp_structure is None:
            logger.info("初始化PP-Structure（表格识别模式）...")
            try:
                _pp_structure = PPStructure(
                    show_log=False,
                    image_orientation=True,
                    table=True,
                    ocr=True,
                    layout=True,
                    recovery=True,
                    use_gpu=False  # 根据需要设置
                )
                logger.info("✅ PP-Structure初始化成功")
            except Exception as e:
                logger.warning(f"PP-Structure初始化失败: {e}，降级到普通 OCR 模式")
                return get_paddle_ocr(use_structure=False)
        return _pp_structure
    else:
        if _paddle_ocr is None:
            logger.info("初始化PaddleOCR...")
            try:
                _paddle_ocr = PaddleOCR(
                    use_angle_cls=True,
                    lang='ch',
                    use_gpu=False,
                    show_log=False
                )
                logger.info("✅ PaddleOCR初始化成功")
            except Exception as e:
                logger.error(f"PaddleOCR初始化失败: {e}")
                raise
        return _paddle_ocr

# ========== 文档分析函数 ==========
def analyze_pdf_with_paddleocr(file_path: str, use_structure: bool = True) -> Dict:
    """使用PaddleOCR分析PDF"""
    try:
        # 先尝试用pypdf提取文本
        if PYPDF_AVAILABLE:
            reader = pypdf.PdfReader(file_path)
            full_text = ""

            for page_num, page in enumerate(reader.pages):
                text = page.extract_text()
                if text and len(text.strip()) > 100:
                    full_text += f"=== 第 {page_num + 1} 页 ===\n\n"
                    full_text += text + "\n\n"

            if full_text.strip():
                # 成功提取到文本
                page_data = {
                    "page_number": 0,
                    "text": full_text.strip(),
                    "text_length": len(full_text),
                    "layouts": [],
                    "tables": []
                }

                return {
                    "filename": Path(file_path).name,
                    "mime_type": "application/pdf",
                    "pages": [page_data],
                    "summary": {
                        "total_pages": 1,
                        "total_text_length": len(full_text),
                        "table_count": 0,
                        "processing_method": "pypdf"
                    }
                }

        # 如果pypdf没有提取到足够文本，使用PaddleOCR
        logger.info("使用PaddleOCR进行OCR识别...")
        ocr_engine = get_paddle_ocr(use_structure=use_structure)

        # 将PDF转换为图片并识别
        import fitz
        import io
        from PIL import Image

        doc = fitz.open(file_path)
        pages_data = []

        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            mat = fitz.Matrix(2, 2)
            pix = page.get_pixmap(matrix=mat)
            img_bytes = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_bytes))

            if use_structure:
                result = ocr_engine(img)
                # 解析PP-Structure结果
                page_data = parse_structure_result(result, page_num + 1)
            else:
                result = ocr_engine.ocr(img)
                page_data = parse_ocr_result(result, page_num + 1)

            pages_data.append(page_data)

        doc.close()

        return {
            "filename": Path(file_path).name,
            "mime_type": "application/pdf",
            "pages": pages_data,
            "summary": {
                "total_pages": len(pages_data),
                "total_text_length": sum(p.get("text_length", 0) for p in pages_data),
                "table_count": sum(len(p.get("tables", [])) for p in pages_data),
                "processing_method": "paddleocr"
            }
        }

    except Exception as e:
        logger.error(f"PDF分析失败: {e}")
        raise

def parse_structure_result(result, page_num: int) -> Dict:
    """解析PP-Structure结果"""
    page_data = {
        "page_number": page_num,
        "text": "",
        "text_length": 0,
        "layouts": [],
        "tables": []
    }

    text_parts = []
    tables_found = 0

    for item in result:
        if item['type'] == 'text':
            res = item['res']
            if isinstance(res, list):
                text = "\n".join([line[0] for line in res])
            else:
                text = str(res)
            text_parts.append(text)
        elif item['type'] == 'table':
            tables_found += 1
            table_data = item['res']
            html = table_data.get('html', '')

            # 解析HTML表格
            table_info = parse_html_table(html)
            table_info['bbox'] = item.get('bbox', [])
            page_data["tables"].append(table_info)

    page_data["text"] = "\n\n".join(text_parts)
    page_data["text_length"] = len(page_data["text"])

    return page_data

def parse_ocr_result(result, page_num: int) -> Dict:
    """解析普通OCR结果"""
    page_data = {
        "page_number": page_num,
        "text": "",
        "text_length": 0,
        "layouts": [],
        "tables": []
    }

    text_lines = []
    for page in result:
        for line in page:
            if line:
                text_lines.append(line[0])

    page_data["text"] = "\n".join(text_lines)
    page_data["text_length"] = len(page_data["text"])

    return page_data

def parse_html_table(html: str) -> Dict:
    """解析HTML表格"""
    try:
        dfs = pd.read_html(html)
        if not dfs:
            return {"rows": 0, "columns": 0, "data": []}

        df = dfs[0]
        df = df.fillna('')

        return {
            "rows": len(df),
            "columns": len(df.columns),
            "column_names": list(df.columns),
            "data": df.to_dict(orient='records')
        }
    except Exception as e:
        logger.warning(f"HTML表格解析失败: {e}")
        return {"rows": 0, "columns": 0, "data": [], "error": str(e)}

def analyze_docx(file_path: str) -> Dict:
    """分析Word文档"""
    if not OFFICE_SUPPORT:
        raise HTTPException(status_code=500, detail="Office支持未启用")

    try:
        doc = DocxDocument(file_path)
        text_content = []
        tables_data = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())

        for table_num, table in enumerate(doc.tables):
            table_rows = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_rows.append(row_data)

            if table_rows:
                columns = table_rows[0]
                data_rows = table_rows[1:]
                df = pd.DataFrame(data_rows, columns=columns)

                tables_data.append({
                    "table_number": table_num + 1,
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": list(df.columns),
                    "data": df.to_dict(orient='records')
                })

        full_text = "\n\n".join(text_content)

        return {
            "filename": Path(file_path).name,
            "mime_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "pages": [{
                "page_number": 1,
                "text": full_text,
                "text_length": len(full_text),
                "layouts": [],
                "tables": tables_data
            }],
            "summary": {
                "total_pages": 1,
                "total_text_length": len(full_text),
                "table_count": len(tables_data)
            }
        }
    except Exception as e:
        logger.error(f"Word文档分析失败: {e}")
        raise

def analyze_xlsx(file_path: str) -> Dict:
    """分析Excel文档"""
    if not OFFICE_SUPPORT:
        raise HTTPException(status_code=500, detail="Office支持未启用")

    try:
        with pd.ExcelFile(file_path) as excel_file:
            text_content = []
            tables_data = []

            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                sheet_text = f"## 工作表: {sheet_name}\n\n"
                sheet_text += df.to_string(index=False, na_rep='')
                text_content.append(sheet_text)

                tables_data.append({
                    "sheet_name": sheet_name,
                    "rows": len(df),
                    "columns": len(df.columns),
                    "column_names": list(df.columns),
                    "data": df.to_dict(orient='records')
                })

        full_text = "\n\n".join(text_content)

        return {
            "filename": Path(file_path).name,
            "mime_type": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "pages": [{
                "page_number": 1,
                "text": full_text,
                "text_length": len(full_text),
                "layouts": [],
                "tables": tables_data
            }],
            "summary": {
                "total_pages": 1,
                "total_text_length": len(full_text),
                "table_count": len(tables_data)
            }
        }
    except Exception as e:
        logger.error(f"Excel分析失败: {e}")
        raise

# ========== 向量检索服务函数 ==========
def get_embedding_service():
    """获取嵌入服务"""
    global _embedding_service
    if _embedding_service is None and VECTOR_FEATURES_ENABLED:
        _embedding_service = EmbeddingService()
        _embedding_service.load_model()
    return _embedding_service

def get_vector_storage():
    """获取向量存储"""
    global _vector_storage
    if _vector_storage is None and VECTOR_FEATURES_ENABLED:
        _vector_storage = VectorStorage()
        _vector_storage.connect()
    return _vector_storage

def get_hybrid_retriever():
    """获取混合检索器"""
    global _hybrid_retriever
    if _hybrid_retriever is None and HYBRID_SEARCH_ENABLED:
        embedding_service = get_embedding_service()
        vector_storage = get_vector_storage()
        _hybrid_retriever = HybridRetriever(embedding_service, vector_storage)
    return _hybrid_retriever

# ========== API端点 ==========

@app.get("/test/service")
async def test_service():
    """测试端点 - 检查服务实例状态"""
    return {
        "service_instance": str(type(service_instance.get_chat_service())) if service_instance.get_chat_service() else None,
        "sys_modules": str(type(sys.modules.get('_chat_service_instance'))) if sys.modules.get('_chat_service_instance') else None
    }


@app.get("/health")
async def health_check():
    """健康检查"""
    try:
        paddleocr_ok = PADDLEOCR_AVAILABLE
        vector_ok = VECTOR_FEATURES_ENABLED
        hybrid_ok = HYBRID_SEARCH_ENABLED

        import torch
        overall_status = "healthy"

        return {
            "status": overall_status,
            "services": {
                "paddleocr": paddleocr_ok,
                "vector_features": vector_ok,
                "hybrid_search": hybrid_ok
            },
            "cuda_available": torch.cuda.is_available() if VECTOR_FEATURES_ENABLED else False,
            "timestamp": datetime.datetime.now().isoformat()
        }
    except Exception as e:
        logger.error(f"健康检查失败: {e}")
        return {"status": "error", "error": str(e)}

@app.post("/analyze")
async def analyze_document(file: UploadFile = File(...)):
    """分析PDF文档"""
    allowed_mime_types = {
        "application/pdf": "pdf"
    }

    if file.content_type not in allowed_mime_types:
        raise HTTPException(status_code=400, detail="只支持PDF文件")

    temp_file = None
    try:
        contents = await file.read()
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(contents)
            temp_file = tmp.name

        result = analyze_pdf_with_paddleocr(temp_file, use_structure=True)

        return {
            "status": "success",
            "data": result
        }

    except Exception as e:
        logger.error(f"文档分析失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if temp_file and os.path.exists(temp_file):
            os.unlink(temp_file)

@app.post("/analyze/office")
async def analyze_office_document(file: UploadFile = File(...)):
    """分析Office文档"""
    allowed_mime_types = {
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx"
    }

    if file.content_type not in allowed_mime_types:
        raise HTTPException(status_code=400, detail="不支持的文件类型")

    temp_file = None
    try:
        contents = await file.read()
        file_ext = allowed_mime_types[file.content_type]

        with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_ext}") as tmp:
            tmp.write(contents)
            temp_file = tmp.name

        if file_ext == "docx":
            result = analyze_docx(temp_file)
        elif file_ext == "xlsx":
            result = analyze_xlsx(temp_file)

        return {
            "status": "success",
            "data": result
        }

    except Exception as e:
        logger.error(f"Office文档分析失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if temp_file and os.path.exists(temp_file):
            os.unlink(temp_file)

@app.post("/ingest")
async def ingest_document(file: UploadFile = File(...)):
    """摄取文档到向量数据库"""
    allowed_mime_types = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
        "application/vnd.ms-excel": "xls"
    }

    # 扩展名到 MIME 类型的映射（用于 content_type 为 None 的情况）
    ext_to_mime = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        ".xls": "application/vnd.ms-excel"
    }

    temp_file = None
    doc_id = str(uuid.uuid4())

    try:
        contents = await file.read()

        # 检查文件类型 - 支持 MIME 类型或文件扩展名
        file_ext = None

        # 首先尝试从 content_type 获取
        if file.content_type in allowed_mime_types:
            file_ext = allowed_mime_types[file.content_type]
        else:
            # 如果 content_type 为 None 或不匹配，从文件名获取扩展名
            if file.filename:
                import os
                _, ext = os.path.splitext(file.filename)
                ext = ext.lower()
                if ext == '.pdf':
                    file_ext = 'pdf'
                elif ext == '.docx':
                    file_ext = 'docx'
                elif ext == '.xlsx' or ext == '.xls':
                    file_ext = ext[1:]  # 去掉点号

        if not file_ext:
            logger.error(f"不支持的文件类型: content_type={file.content_type}, filename={file.filename}")
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型。支持的类型: PDF, DOCX, XLSX, XLS"
            )

        logger.info(f"处理文件: {file.filename}, content_type: {file.content_type}, 检测到的扩展名: {file_ext}")

        # 使用mkstemp创建临时文件，避免文件句柄问题
        fd, temp_file = tempfile.mkstemp(suffix=f".{file_ext}")
        try:
            with os.fdopen(fd, 'wb') as tmp:
                tmp.write(contents)
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"文件写入失败: {str(e)}")
        
        # 分析文档提取文本
        logger.info(f"开始分析文档: {temp_file}")
        if file_ext == "pdf":
            analysis_result = analyze_pdf_with_paddleocr(temp_file)
        elif file_ext == "docx":
            analysis_result = analyze_docx(temp_file)
        elif file_ext == "xlsx":
            analysis_result = analyze_xlsx(temp_file)
        else:
            raise HTTPException(status_code=400, detail=f"不支持的文件扩展名: {file_ext}")

        if not analysis_result:
            raise HTTPException(status_code=500, detail="文档分析失败，未返回结果")

        # 提取文本
        full_text = ""
        for page in analysis_result.get("pages", []):
            full_text += page.get("text", "") + "\n\n"

        logger.info(f"提取的文本长度: {len(full_text)} 字符")

        full_text = full_text.strip()

        if not full_text:
            raise HTTPException(status_code=400, detail="文档没有提取到文本内容")

        # 向量化
        vector_storage = get_vector_storage()
        if not vector_storage:
            raise HTTPException(status_code=500, detail="向量存储服务未初始化")

        chunks = vector_storage.chunk_text(full_text)

        if not chunks:
            raise HTTPException(status_code=400, detail="文档分块失败")

        embedding_service = get_embedding_service()
        if not embedding_service:
            raise HTTPException(status_code=500, detail="嵌入服务未初始化")

        embeddings = embedding_service.embed_batch(chunks)

        metadata = {
            "doc_id": doc_id,
            "filename": file.filename,
            "file_size": len(contents),
            "mime_type": file.content_type,
            "ingested_at": datetime.datetime.now().isoformat(),
            "total_chunks": len(chunks)
        }

        storage_result = vector_storage.store_document(doc_id, chunks, embeddings, metadata)

        if not storage_result.get("success"):
            error_msg = storage_result.get("error", "未知存储错误")
            logger.error(f"存储失败: {error_msg}")
            raise HTTPException(status_code=500, detail=f"存储失败: {error_msg}")

        return {
            "status": "success",
            "data": {
                "doc_id": doc_id,
                "filename": file.filename,
                "total_chunks": len(chunks),
                "stored_chunks": storage_result.get("stored_chunks", 0)
            }
        }

    except HTTPException:
        raise
    except Exception as e:
        import traceback
        error_detail = str(e) if str(e) else f"{type(e).__name__}: 未知错误"
        logger.error(f"文档摄取失败: {error_detail}")
        logger.error(f"异常详情: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=error_detail)
    finally:
            # 清理临时文件
            if temp_file and os.path.exists(temp_file):
                try:
                    os.unlink(temp_file)
                except Exception:
                    # 忽略删除错误
                    pass

@app.post("/search")
async def search_documents(
    query: str,
    top_k: int = Query(10, ge=1, le=100),
    score_threshold: float = Query(0.3, ge=0, le=1)
):
    """向量搜索"""
    if not query or len(query.strip()) < 2:
        raise HTTPException(status_code=400, detail="查询太短")

    try:
        embedding_service = get_embedding_service()
        if not embedding_service:
            raise HTTPException(status_code=500, detail="嵌入服务未初始化")

        query_vector = embedding_service.embed_text(query)

        vector_storage = get_vector_storage()
        if not vector_storage:
            raise HTTPException(status_code=500, detail="向量存储服务未初始化")

        results = vector_storage.search(
            query_vector=query_vector,
            top_k=top_k,
            score_threshold=score_threshold
        )

        formatted_results = []
        for result in results:
            formatted_results.append({
                "id": result.get("id", ""),
                "text": result.get("text", "")[:500],
                "score": round(result.get("score", 0), 4),
                "source_document": result.get("source_document", "unknown")
            })

        return {
            "status": "success",
            "query": query,
            "total_results": len(results),
            "results": formatted_results
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"搜索失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/search/hybrid")
async def hybrid_search_documents(
    query: str,
    top_k: int = Query(10, ge=1, le=50),
    fusion_method: str = Query("rrf", pattern="^(rrf|weighted|simple)$"),
    vector_weight: float = Query(0.7, ge=0, le=1),
    keyword_weight: float = Query(0.3, ge=0, le=1)
):
    """混合搜索"""
    if not HYBRID_SEARCH_ENABLED:
        raise HTTPException(status_code=501, detail="混合搜索功能未启用")

    if not query or len(query.strip()) < 2:
        raise HTTPException(status_code=400, detail="查询太短")

    try:
        retriever = get_hybrid_retriever()
        if not retriever:
            raise HTTPException(status_code=500, detail="混合检索器未初始化")

        results = retriever.hybrid_search(
            query=query,
            top_k=top_k,
            fusion_method=fusion_method,
            vector_weight=vector_weight,
            keyword_weight=keyword_weight
        )

        formatted_results = []
        for result in results:
            formatted_result = {
                "id": result.get("id", ""),
                "text": result.get("text", "")[:500],
                "score": round(result.get("fusion_score", result.get("score", 0)), 4),
                "source_document": result.get("source_document", "unknown")
            }
            formatted_results.append(formatted_result)

        return {
            "status": "success",
            "query": query,
            "total_results": len(results),
            "results": formatted_results
        }

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"混合搜索失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/api/kb/status")
async def get_kb_status():
    """获取知识库状态"""
    try:
        vector_storage = get_vector_storage()
        if not vector_storage:
            return {
                "success": True,
                "data": {
                    "totalDocuments": 0,
                    "totalChunks": 0,
                    "vectorStatus": "error",
                    "lastUpdated": datetime.datetime.now().isoformat()
                }
            }

        # 获取文档列表统计实际文档数量
        documents = vector_storage.list_all_documents()
        documents_count = len(documents)

        collection_info = vector_storage.get_collection_info()
        chunks_count = collection_info.get("points_count", 0)

        # 判断状态
        vector_status = "healthy"
        if not vector_storage.connected:
            vector_status = "error"
        elif documents_count == 0:
            vector_status = "warning"

        return {
            "success": True,
            "data": {
                "totalDocuments": documents_count,
                "totalChunks": chunks_count,
                "vectorStatus": vector_status,
                "lastUpdated": datetime.datetime.now().isoformat()
            }
        }
    except Exception as e:
        logger.error(f"获取知识库状态失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/kb/reset")
async def reset_kb():
    """重置知识库（清空所有文档）"""
    try:
        vector_storage = get_vector_storage()
        if not vector_storage or not vector_storage.connected:
            return {
                "success": False,
                "error": "向量数据库未连接"
            }

        result = vector_storage.clear_collection()

        if result.get("success"):
            cleared_points = result.get("cleared_points", 0)
            # 确保只返回简单类型，避免Pydantic验证错误
            return {
                "success": True,
                "message": f"知识库已重置，清空了 {cleared_points} 个文档块",
                "data": {
                    "cleared_points": int(cleared_points) if cleared_points is not None else 0
                }
            }
        else:
            logger.error(f"❌ 知识库重置失败: {result.get('error')}")
            return {
                "success": False,
                "error": str(result.get("error", "未知错误"))
            }
    except Exception as e:
        logger.error(f"重置知识库失败: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ========== LangGraph多Agent系统（双模型版本）==========

def get_active_chat_service():
    """
    获取当前激活的聊天服务

    优先级: 双模型服务 > 单模型服务(AWQ) > 旧版单模型服务(Ollama)

    Returns:
        当前激活的聊天服务实例
    """
    import sys

    # 尝试从 __main__ 模块获取（运行 python main.py 时的模块名）
    main_module = sys.modules.get('__main__') or sys.modules.get('main')
    if main_module:
        dual_model = getattr(main_module, '_dual_model_service', None)
        single_model = getattr(main_module, '_single_model_service', None)
        chat_service = getattr(main_module, '_chat_service', None)

        # 按优先级返回
        if dual_model is not None:
            return dual_model
        if single_model is not None:
            return single_model
        if chat_service is not None:
            return chat_service

    # 如果都没找到，返回 None（会由调用方处理）
    return None


@app.get("/debug/services")
async def debug_services():
    """调试端点 - 检查服务状态"""
    import sys

    main_module = sys.modules.get('__main__') or sys.modules.get('main')
    if main_module:
        dual_model = getattr(main_module, '_dual_model_service', None)
        single_model = getattr(main_module, '_single_model_service', None)
        chat_service = getattr(main_module, '_chat_service', None)

        return {
            "_dual_model_service": str(type(dual_model)) if dual_model else None,
            "_single_model_service": str(type(single_model)) if single_model else None,
            "_chat_service": str(type(chat_service)) if chat_service else None,
            "dual_model_service_is_none": dual_model is None,
            "single_model_service_is_none": single_model is None,
            "chat_service_is_none": chat_service is None,
            "main_module": str(main_module),
            "main_module_name": main_module.__name__
        }

    return {"error": "Could not find main module"}


# ========== API端点 ==========

# 会话历史 API (兼容前端 /api/conversations 路径)
@app.get("/api/conversations")
async def get_conversations_list():
    """获取会话历史列表"""
    from routes.chat_routes import get_conversations
    return await get_conversations()


@app.get("/api/conversations/{conversation_id}/messages")
async def get_conversation_messages_detail(conversation_id: str):
    """获取会话的消息历史"""
    from routes.chat_routes import get_conversation_messages
    return await get_conversation_messages(conversation_id)


# 批量删除路由必须放在参数化路由之前
@app.delete("/api/conversations/batch")
async def batch_delete_conversations_endpoint(http_request: Request):
    """批量删除会话"""
    logger.info("=== 批量删除端点被调用 ===")
    try:
        # 从请求体获取数据
        body = await http_request.json()
        ids = body.get("ids", [])
        logger.info(f"收到的IDs: {ids}")

        if not ids:
            return {
                "success": False,
                "error": "缺少ids参数"
            }

        from routes.mobile_routes import _sessions_storage, _user_sessions
        logger.info(f"会话存储中有 {len(_sessions_storage)} 个会话")

        deleted_count = 0
        failed_ids = []

        for conversation_id in ids:
            try:
                # 检查会话是否存在
                if conversation_id not in _sessions_storage:
                    failed_ids.append(conversation_id)
                    continue

                # 获取会话的user_id（从会话数据中获取）
                session = _sessions_storage[conversation_id]
                user_id = session.get("user_id", "default_user")

                # 从用户索引中移除
                if user_id in _user_sessions:
                    if conversation_id in _user_sessions[user_id]:
                        _user_sessions[user_id].remove(conversation_id)

                # 删除会话
                del _sessions_storage[conversation_id]
                deleted_count += 1

            except Exception as e:
                logger.warning(f"删除会话 {conversation_id} 失败: {e}")
                failed_ids.append(conversation_id)

        # 如果全部失败，返回失败
        if deleted_count == 0:
            return {
                "success": False,
                "error": "没有删除任何会话",
                "deleted_count": 0,
                "failed_ids": failed_ids
            }

        return {
            "success": True,
            "message": f"已删除{deleted_count}个会话",
            "deleted_count": deleted_count,
            "failed_ids": failed_ids
        }

    except Exception as e:
        logger.error(f"批量删除会话失败: {e}")
        return {
            "success": False,
            "error": str(e),
            "deleted_count": 0,
            "failed_ids": ids if 'ids' in locals() else []
        }


@app.get("/api/conversations/debug")
async def debug_sessions():
    """调试端点 - 查看会话存储状态"""
    from routes.mobile_routes import _sessions_storage, _user_sessions
    return {
        "success": True,
        "sessions_count": len(_sessions_storage),
        "sessions_ids": list(_sessions_storage.keys()),
        "user_sessions": _user_sessions.get("default_user", [])
    }


@app.delete("/api/conversations/{conversation_id}")
async def delete_conversation_endpoint(conversation_id: str):
    """删除会话"""
    from routes.chat_routes import delete_conversation
    return await delete_conversation(conversation_id)


# 注册新路由
try:
    from routes.chat_routes import router as chat_router
    from routes.document_routes import router as document_router
    from routes.mobile_routes import router as mobile_router

    app.include_router(chat_router)
    app.include_router(document_router)
    app.include_router(mobile_router)

    logger.info("✅ 新路由已注册: /api/chat, /api/documents, /api/mobile")

except ImportError as e:
    logger.warning(f"⚠️  无法导入新路由: {e}")


if __name__ == "__main__":
    import argparse
    import signal

    parser = argparse.ArgumentParser(description="DeepDoc API Server")
    parser.add_argument("--port", type=int, default=8000, help="Port to run the server on")
    parser.add_argument("--host", type=str, default="0.0.0.0", help="Host to bind to")
    parser.add_argument("--timeout-keep-alive", type=int, default=30, help="Keep-alive timeout in seconds")
    parser.add_argument("--timeout-graceful-shutdown", type=int, default=10, help="Graceful shutdown timeout in seconds")
    args = parser.parse_args()

    # 配置 uvicorn
    config = uvicorn.Config(
        app,
        host=args.host,
        port=args.port,
        timeout_keep_alive=args.timeout_keep_alive,
        timeout_graceful_shutdown=args.timeout_graceful_shutdown,
        log_level="info",
        access_log=True
    )

    server = uvicorn.Server(config)

    # 信号处理器
    def handle_signal(signum, frame):
        logger.info(f"📴 收到信号 {signum}，准备关闭服务器...")
        # 通知 uvicorn 服务器应该关闭
        if server.should_exit:
            logger.warning("⚠️  强制退出服务器...")
            import os
            os._exit(1)
        else:
            server.should_exit = True

    # 注册信号处理器
    signal.signal(signal.SIGINT, handle_signal)
    signal.signal(signal.SIGTERM, handle_signal)

    logger.info(f"🚀 启动服务器: http://{args.host}:{args.port}")
    logger.info(f"⏱️  超时配置: keep-alive={args.timeout_keep_alive}s, graceful-shutdown={args.timeout_graceful_shutdown}s")

    try:
        server.run()
    except KeyboardInterrupt:
        logger.info("👋 用户中断，正在关闭...")
    except Exception as e:
        logger.error(f"❌ 服务器错误: {e}")
        import traceback
        logger.error(traceback.format_exc())
    finally:
        logger.info("✅ 服务器已停止")
