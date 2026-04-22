"""极简长上下文服务 - 完整版
支持文档加载、表格解析、关键词搜索、DeepSeek API集成
"""
import logging
import re
import os
from pathlib import Path
from typing import Dict, List, Optional, Tuple
import requests
import json

logger = logging.getLogger(__name__)


class SimpleLongContextService:
    """极简长上下文服务 - 核心实现"""

    def __init__(self, deepseek_api_key: str = None):
        self.documents: Dict[str, str] = {}
        self.document_metadata: Dict[str, dict] = {}  # 存储文档元数据
        self.deepseek_api_key = deepseek_api_key or os.getenv("DEEPSEEK_API_KEY")
        self.deepseek_api_url = "https://api.deepseek.com/v1/chat/completions"

        # 系统提示词
        self.system_prompt = self._load_system_prompt()

        logger.info("🚀 极简长上下文服务初始化完成")

    def _load_system_prompt(self) -> str:
        """加载系统提示词"""
        try:
            from config.system_prompt import get_system_prompt
            return get_system_prompt()
        except ImportError:
            # 回退到硬编码提示词
            return """你是运营商渠道业务AI助手，专门查询运营商政策文档。

## 核心原则
1. **严格依据文档回答**：只能基于提供的文档内容回答，不能编造信息
2. **简洁直接**：找到答案直接回答，找不到直接说找不到，不要啰嗦
3. **准确优先**：不确定的信息不要提供

## 输出格式要求（严格执行）

【回答】
（直接回答，最多3-5句话，禁止废话）

**如果找到答案：**
直接给出关键信息，用1.2.3.组织要点

**如果找不到答案：**
直接说明"当前文档未找到相关产品ID和资费信息"，不要提供无关信息

【数据来源】
（列出引用的文档名称）

## 重要约束
- 总字数不超过120字
- 禁止使用markdown格式
- 禁止提供推测或无关信息
- 找不到答案时，一句话说明即可

严格按照此格式，保持简洁专业。"""

    def load_documents(self, data_dir: str) -> int:
        """从目录加载所有文档（DOCX, XLSX）"""
        data_path = Path(data_dir)
        if not data_path.exists():
            logger.warning(f"⚠️ 文档目录不存在: {data_path}")
            return 0

        count = 0
        total_size = 0

        # 加载 DOCX
        for file_path in data_path.rglob("*.docx"):
            try:
                content, metadata = self._parse_docx_enhanced(file_path)
                if content:
                    self.documents[file_path.name] = content
                    self.document_metadata[file_path.name] = metadata
                    count += 1
                    total_size += file_path.stat().st_size
                    logger.info(f"✅ 加载 DOCX: {file_path.name} ({metadata['word_count']} 字, {len(metadata['tables'])} 表格)")
            except Exception as e:
                logger.error(f"❌ DOCX 加载失败 {file_path.name}: {e}")

        # 加载 XLSX
        for file_path in data_path.rglob("*.xlsx"):
            try:
                content, metadata = self._parse_xlsx_enhanced(file_path)
                if content:
                    self.documents[file_path.name] = content
                    self.document_metadata[file_path.name] = metadata
                    count += 1
                    total_size += file_path.stat().st_size
                    logger.info(f"✅ 加载 XLSX: {file_path.name} ({metadata['sheet_count']} 工作表, {metadata['total_rows']} 行)")
            except Exception as e:
                logger.error(f"❌ XLSX 加载失败 {file_path.name}: {e}")

        logger.info(f"📚 文档加载完成: {count} 个文档, 总计 {self._format_size(total_size)}")
        return count

    def _parse_docx_enhanced(self, file_path: Path) -> Tuple[str, dict]:
        """增强版 DOCX 解析 - 完整表格处理"""
        try:
            from docx import Document
            import pandas as pd

            doc = Document(str(file_path))
            content_parts = []
            tables_info = []
            word_count = 0

            # 遍历文档元素，保持顺序
            for element in doc.element.body:
                if element.tag.endswith('p'):
                    # 段落处理
                    for para in doc.paragraphs:
                        if para._element == element:
                            text = para.text.strip()
                            if text:
                                content_parts.append(text)
                                word_count += len(text)
                            break
                elif element.tag.endswith('tbl'):
                    # 表格处理 - 完整版
                    for table in doc.tables:
                        if table._element == element:
                            table_content = self._parse_table_enhanced(table)
                            if table_content:
                                content_parts.append(table_content)
                                tables_info.append({
                                    'rows': len(table.rows),
                                    'cols': len(table.columns),
                                    'content': table_content  # 保存完整表格内容
                                })
                            break

            result = "\n\n".join(content_parts)
            metadata = {
                'word_count': word_count,
                'tables': tables_info,
                'paragraph_count': len([p for p in doc.paragraphs if p.text.strip()])
            }
            return result, metadata

        except Exception as e:
            logger.error(f"❌ DOCX 解析失败: {e}")
            return self._fallback_docx_text(file_path), {'word_count': 0, 'tables': []}

    def _parse_table_enhanced(self, table) -> str:
        """增强版表格解析 - 完整保留表格信息"""
        try:
            import pandas as pd

            # 提取表格数据
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    # 处理单元格中的换行，但保留重要结构
                    cell_text = cell_text.replace('\n', ' ')
                    row_data.append(cell_text)
                table_data.append(row_data)

            if not table_data or len(table_data) < 1:
                return ""

            # 创建完整的文本表格
            table_lines = []
            for i, row in enumerate(table_data):
                if i == 0:
                    # 表头
                    table_lines.append(" | ".join(row))
                    table_lines.append("-+-".join(["-" for _ in row]))
                else:
                    # 数据行
                    table_lines.append(" | ".join(row))

            return "\n" + "\n".join(table_lines) + "\n"

        except Exception as e:
            logger.error(f"❌ 表格解析失败: {e}")
            # 回退：直接用原始数据
            return "\n".join([" | ".join(row) for row in table_data]) + "\n"

    def _parse_xlsx_enhanced(self, file_path: Path) -> Tuple[str, dict]:
        """增强版 XLSX 解析"""
        try:
            import pandas as pd

            excel_file = pd.ExcelFile(str(file_path))
            content_parts = []
            total_rows = 0
            sheet_count = len(excel_file.sheet_names)

            for sheet_name in excel_file.sheet_names:
                try:
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    total_rows += len(df)

                    # 转换为 Markdown 表格
                    markdown_table = df.to_markdown(index=False)
                    sheet_content = f"## 工作表: {sheet_name}\n\n{markdown_table}"
                    content_parts.append(sheet_content)

                except Exception as e:
                    # 回退到字符串格式
                    df = pd.read_excel(excel_file, sheet_name=sheet_name)
                    total_rows += len(df)
                    sheet_content = f"## 工作表: {sheet_name}\n\n{df.to_string(index=False)}"
                    content_parts.append(sheet_content)

            result = "\n\n".join(content_parts)
            metadata = {
                'sheet_count': sheet_count,
                'total_rows': total_rows,
                'total_sheets': sheet_count
            }
            return result, metadata

        except Exception as e:
            logger.error(f"❌ XLSX 解析失败: {e}")
            return f"[无法解析 Excel 文件: {file_path.name}]", {'sheet_count': 0, 'total_rows': 0}

    def _fallback_docx_text(self, file_path: Path) -> str:
        """DOCX 回退解析"""
        try:
            from docx import Document
            doc = Document(str(file_path))
            return "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception as e:
            logger.error(f"❌ DOCX 回退解析失败: {e}")
            return f"[无法解析: {file_path.name}]"

    def search_documents(self, query: str, top_k: int = 5) -> List[dict]:
        """增强版文档搜索 - 改进的BM25风格，支持中文分词"""
        if not query.strip():
            return []

        query_lower = query.lower()
        # 停用词
        stop_words = {'的', '了', '是', '在', '有', '和', '与', '及', '等', '或', '但', '而'}

        # 改进的中文分词策略
        query_terms = []

        # 首先尝试按空格分词（适用于西文或已分词文本）
        space_split = [t for t in query_lower.split() if t not in stop_words and len(t) > 1]
        if space_split:
            query_terms.extend(space_split)

        # 对于中文，提取2字以上的词语（简单bigram + 词组提取）
        if self._is_chinese(query_lower):
            # 提取所有连续的中文字符序列作为候选词
            chinese_words = self._extract_chinese_words(query_lower, stop_words)
            query_terms.extend(chinese_words)

            # 如果没有找到合适的词组，使用单个重要字符
            if not query_terms:
                important_chars = [c for c in query_lower if c not in stop_words and len(c) > 0]
                query_terms.extend(important_chars)

        # 去重并过滤
        query_terms = list(set(query_terms))
        query_terms = [t for t in query_terms if len(t) >= 1]

        # 如果仍然没有有效关键词，使用原查询
        if not query_terms:
            query_terms = [query_lower]

        results = []

        for doc_name, content in self.documents.items():
            content_lower = content.lower()
            score = 0
            snippets = []

            # 对每个关键词计算匹配分数
            for term in query_terms:
                matches = content_lower.count(term)
                if matches > 0:
                    score += matches * 5  # 每个匹配加5分

                    # 找到所有出现位置，提取上下文
                    start = 0
                    while True:
                        idx = content_lower.find(term, start)
                        if idx == -1:
                            break

                        # 提取该位置周围的上下文
                        snippet_start = max(0, idx - 150)
                        snippet_end = min(len(content), idx + len(term) + 200)
                        snippet = content[snippet_start:snippet_end].replace("\n", " ")
                        snippets.append({
                            "term": term,
                            "position": idx,
                            "snippet": snippet
                        })
                        start = idx + 1

            if score > 0:
                # 选择最佳的片段
                if snippets:
                    best = max(snippets, key=lambda x: len([t for t in query_terms if t in x["snippet"]]))
                    results.append({
                        "document": doc_name,
                        "score": score,
                        "snippet": best["snippet"],
                        "metadata": self.document_metadata.get(doc_name, {})
                    })

        # 按分数排序
        results.sort(key=lambda x: x["score"], reverse=True)
        return results[:top_k]

    def ask_deepseek_long_context(self, question: str, conversation_history: list = None) -> dict:
        """调用 DeepSeek API 进行问答 - 真正的长上下文架构

        Args:
            question: 用户问题
            conversation_history: 对话历史列表，格式为 [{"role": "user/assistant", "content": "..."}]
        """
        if not self.deepseek_api_key:
            logger.warning("⚠️ DeepSeek API Key 未配置")
            return self._fallback_answer_long_context(question)

        try:
            # 直接构造所有文档的完整内容
            all_docs_content = self._build_all_documents_context()

            # 构造消息 - 支持对话历史
            messages = [
                {"role": "system", "content": self.system_prompt},
            ]

            # 添加对话历史（最近的对话，最多保留最近6轮以避免token超限）
            if conversation_history and len(conversation_history) > 0:
                logger.info(f"📝 对话历史: {len(conversation_history)} 条消息")
                # 只取最近的对话历史，最多6轮（12条消息）
                recent_history = conversation_history[-12:] if len(conversation_history) > 12 else conversation_history
                for msg in recent_history:
                    # 只添加role和content，避免添加timestamp等额外字段
                    if "role" in msg and "content" in msg:
                        messages.append({
                            "role": msg["role"],
                            "content": msg["content"]
                        })
                        logger.info(f"  - {msg['role']}: {msg['content'][:50]}...")
            else:
                logger.info("📝 无对话历史（首次提问）")

            # 添加当前问题，附带文档上下文和智能上下文提示
            context_hint = self._generate_context_hint(conversation_history, question)

            # 如果有上下文提示，修改用户问题使其更加明确
            enhanced_question = question
            if context_hint:
                # 如果问题是关于"该服务"的产品ID/资费，且上下文提示指明了增值业务
                if "该服务" in question and ("产品ID" in question or "资费" in question):
                    # 从上下文提示中提取服务名称并直接替换问题
                    import re
                    service_match = re.search(r'增值业务[：:](.+?)[。，]', context_hint)
                    if service_match:
                        service_name = service_match.group(1).strip()
                        # 使用英文增强版本，因为DeepSeek对英文问题理解更准确
                        enhanced_question = f"What is the product ID and price for the service mentioned above: {service_name}?"
                        logger.info(f"🔄 问题增强: {question} -> {enhanced_question}")

            user_content = f"参考文档内容：\n\n{all_docs_content}\n\n用户问题：{enhanced_question}\n\n请基于上述文档内容回答问题，并注明引用来源。"

            if context_hint:
                user_content += f"\n\n【上下文提示】{context_hint}"

            messages.append({
                "role": "user",
                "content": user_content
            })

            # 调用 API
            headers = {
                "Content-Type": "application/json; charset=utf-8",
                "Authorization": f"Bearer {self.deepseek_api_key}"
            }

            payload = {
                "model": "deepseek-chat",
                "messages": messages,
                "temperature": 0.1,  # 降低随机性，减少啰嗦
                "max_tokens": 800     # 限制输出长度，强制简洁
            }

            logger.info(f"🔍 调用DeepSeek API (长上下文): question='{question[:50]}...', total_docs={len(self.documents)}")

            response = requests.post(
                self.deepseek_api_url,
                headers=headers,
                json=payload,
                timeout=30
            )

            if response.status_code == 200:
                result = response.json()
                raw_answer = result["choices"][0]["message"]["content"]
                logger.info(f"✅ DeepSeek API 返回成功: answer_length={len(raw_answer)} 字")

                # 后处理：格式化回答为标准格式
                formatted_answer = self._format_answer(raw_answer)

                # 构造符合前端期待的数据格式
                sources_list = []
                for doc_name in self.documents.keys():
                    sources_list.append({
                        "id": doc_name,
                        "filename": doc_name,
                        "relevance": 1.0
                    })

                return {
                    "success": True,
                    "answer": formatted_answer,
                    "sources": sources_list,
                    "model_used": "deepseek-chat-long-context",
                    "context_docs_count": len(self.documents),
                    "all_docs_used": True
                }
            else:
                logger.error(f"❌ DeepSeek API 错误: {response.status_code} - {response.text}")
                return self._fallback_answer_long_context(question)

        except Exception as e:
            logger.error(f"❌ DeepSeek API 调用失败: {e}")
            return self._fallback_answer_long_context(question)

    def _build_all_documents_context(self) -> str:
        """构造所有文档的完整上下文文本"""
        if not self.documents:
            return ""

        context_parts = []
        for doc_name, content in self.documents.items():
            metadata = self.document_metadata.get(doc_name, {})
            word_count = metadata.get('word_count', 0)
            table_count = len(metadata.get('tables', []))

            context_parts.append(
                f"## 文档：{doc_name}\n"
                f"字数：{word_count}，表格数：{table_count}\n"
                f"内容：\n{content}\n"
            )

        return "\n\n".join(context_parts)

    def _fallback_answer_long_context(self, question: str) -> dict:
        """回退答案 - 长上下文架构"""
        if self.documents:
            # 直接显示所有文档的前面部分
            answer_parts = []
            for doc_name, content in list(self.documents.items())[:3]:  # 最多显示3个文档预览
                preview = content[:300].replace('\n', ' ')
                answer_parts.append(
                    f"📄 **{doc_name}**\n"
                    f"内容预览: {preview}...\n"
                )

            answer = (
                f"基于知识库中的 {len(self.documents)} 个文档：\n\n"
                + "\n\n".join(answer_parts) +
                "\n\n💡 *由于未配置 DeepSeek API，以上为文档预览。建议配置 API Key 以获得更准确的答案。*"
            )

            # 构造符合前端期待的数据格式
            sources_list = []
            for doc_name in self.documents.keys():
                sources_list.append({
                    "id": doc_name,
                    "filename": doc_name,
                    "relevance": 1.0
                })

            return {
                "success": True,
                "answer": answer,
                "sources": sources_list,
                "model_used": "direct_document_preview",
                "context_docs_count": len(self.documents)
            }
        else:
            # 构造符合前端期待的数据格式
            sources_list = []
            for doc_name in self.documents.keys():
                sources_list.append({
                    "id": doc_name,
                    "filename": doc_name,
                    "relevance": 1.0
                })

            return {
                "success": True,
                "answer": "未在知识库中找到相关信息，请先上传相关文档。",
                "sources": sources_list,
                "model_used": "no_documents",
                "context_docs_count": 0
            }

    def _is_chinese(self, text: str) -> bool:
        """检查文本是否包含中文"""
        return any('\u4e00' <= char <= '\u9fff' for char in text)

    def _extract_chinese_words(self, text: str, stop_words: set) -> List[str]:
        """从中文文本中提取有意义的词组"""
        words = []

        # 提取2字词组（bigrams）
        for i in range(len(text) - 1):
            word = text[i:i+2]
            if word not in stop_words and '\u4e00' <= word[0] <= '\u9fff':
                words.append(word)

        # 提取3字词组
        for i in range(len(text) - 2):
            word = text[i:i+3]
            if word not in stop_words and '\u4e00' <= word[0] <= '\u9fff':
                words.append(word)

        # 提取4字词组
        for i in range(len(text) - 3):
            word = text[i:i+4]
            if word not in stop_words and '\u4e00' <= word[0] <= '\u9fff':
                words.append(word)

        return words

    def _format_answer(self, raw_answer: str) -> str:
        """格式化回答为标准格式：去除markdown，提取关键信息，重新组织"""
        import re

        # 首先去除已有的【回答】标记（如果存在），准备重新组织
        clean_answer = raw_answer.replace('【回答】', '').strip()

        # 去除markdown格式
        clean_answer = re.sub(r'\*\*(.*?)\*\*', r'\1', clean_answer)  # 去除粗体
        clean_answer = re.sub(r'\*', '', clean_answer)  # 去除列表符号
        clean_answer = re.sub(r'- ', '', clean_answer)  # 去除列表符号
        clean_answer = re.sub(r'###*', '', clean_answer)  # 去除标题
        clean_answer = re.sub(r'\[.*?\]', '', clean_answer)  # 去除引用

        # 如果内容已经包含【数据来源】，需要替换为我们自己的文档列表
        if '【数据来源】' in clean_answer:
            # 分离回答部分和数据来源部分
            parts = clean_answer.split('【数据来源】')
            answer_part = parts[0].strip()
            # 去除markdown格式
            answer_part = re.sub(r'\*\*(.*?)\*\*', r'\1', answer_part)
            answer_part = re.sub(r'###*', '', answer_part)
            # 重新构建格式，使用实际的文档列表
            formatted = f"【回答】\n{answer_part}\n\n【数据来源】\n"
            for doc_name in list(self.documents.keys())[:3]:  # 最多显示3个文档
                formatted += f"{doc_name}\n"
            return formatted

    def _generate_context_hint(self, conversation_history: list, current_question: str) -> str:
        """生成智能上下文提示，帮助AI理解指代关系

        Args:
            conversation_history: 对话历史
            current_question: 当前用户问题

        Returns:
            上下文提示字符串，如果不需要则返回空字符串
        """
        if not conversation_history or len(conversation_history) == 0:
            return ""

        # 检测当前问题是否包含指代词和产品ID/资费相关询问
        reference_words = ["该服务", "这个服务", "上述服务", "该产品", "这个产品", "上述产品", "该业务", "这个业务"]
        id_price_keywords = ["产品ID", "产品id", "资费", "价格", "多少钱"]
        has_reference = any(word in current_question for word in reference_words)
        asks_id_price = any(word in current_question for word in id_price_keywords)

        if not (has_reference and asks_id_price):
            return ""

        # 分析最后一条助手回复，查找增值业务/附加服务
        last_assistant_msg = None
        for msg in reversed(conversation_history):
            if msg.get("role") == "assistant":
                last_assistant_msg = msg.get("content", "")
                break

        if not last_assistant_msg:
            return ""

        # 检测是否提到了增值业务/附加服务 - 使用更精确的模式
        import re
        found_services = []

        # 匹配"自动同订XXX服务"模式
        auto_order_pattern = r'自动同订(.+?服务)'
        auto_matches = re.findall(auto_order_pattern, last_assistant_msg)
        found_services.extend(auto_matches)

        # 匹配"综合意外保障服务"
        if '综合意外保障服务' in last_assistant_msg:
            # 提取完整的服务名称
            accident_pattern = r'(综合意外保障服务（[^）]+）)'
            accident_matches = re.findall(accident_pattern, last_assistant_msg)
            found_services.extend(accident_matches)
            if not accident_matches:
                found_services.append('综合意外保障服务（个人版）')

        # 清理和去重
        clean_services = []
        for service in found_services:
            clean_service = service.strip()
            if 2 < len(clean_service) < 30 and clean_service not in clean_services:
                clean_services.append(clean_service)

        # 如果找到了增值业务，生成明确的上下文提示
        if clean_services:
            # 只取第一个（最相关的）增值业务
            primary_service = clean_services[0]
            return f"增值业务：{primary_service}。请优先提供此增值业务的产品ID和资费信息，不要回答主套餐的信息。"

        return ""

        # 提取关键信息（去除常见的废话）
        noise_phrases = [
            '根据您提供的文档内容',
            '根据当前文档',
            '文档中未找到',
            '建议您',
            '请注意',
            '如果文档信息不完整',
            '具体内容如下',
            '具体规定如下',
            '关于',  # 这个可能在句子中间，需要谨慎处理
            '一、',
            '二、',
            '总结：',
            '数据来源：',
            '引用来源：'
        ]

        lines = clean_answer.split('\n')
        key_lines = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # 跳过包含噪音短语的行（除非有实质性内容）
            if any(noise in line for noise in noise_phrases):
                continue
            key_lines.append(line)

        # 构建简洁的回答
        concise_answer = '\n'.join(key_lines[:10])  # 最多保留10行
        # 进一步简洁化，去除冗余
        concise_answer = re.sub(r'\n{3,}', '\n\n', concise_answer)

        # 添加标准格式
        formatted = f"【回答】\n{concise_answer}\n\n【数据来源】\n"
        for doc_name in list(self.documents.keys())[:3]:  # 最多显示3个文档
            formatted += f"{doc_name}\n"

        return formatted

        # 提取关键信息（去除常见的废话）
        noise_phrases = [
            '根据您提供的文档内容',
            '根据当前文档',
            '文档中未找到',
            '建议您',
            '请注意',
            '如果文档信息不完整',
            '具体内容如下',
            '具体规定如下',
            '关于',  # 这个可能在句子中间，需要谨慎处理
            '一、',
            '二、',
            '总结：',
            '数据来源：',
            '引用来源：'
        ]

        lines = clean_answer.split('\n')
        key_lines = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # 跳过包含噪音短语的行（除非有实质性内容）
            if any(noise in line for noise in noise_phrases):
                continue
            key_lines.append(line)

        # 构建简洁的回答
        concise_answer = '\n'.join(key_lines[:10])  # 最多保留10行
        # 进一步简洁化，去除冗余
        concise_answer = re.sub(r'\n{3,}', '\n\n', concise_answer)

        # 添加标准格式
        formatted = f"【回答】\n{concise_answer}\n\n【数据来源】\n"
        for doc_name in list(self.documents.keys())[:3]:  # 最多显示3个文档
            formatted += f"{doc_name}\n"

        return formatted

    def _format_size(self, size_bytes: int) -> str:
        """格式化文件大小"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f}{unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f}TB"

    def add_document(self, filename: str, content: str, metadata: dict = None) -> None:
        """添加文档到内存"""
        self.documents[filename] = content
        self.document_metadata[filename] = metadata or {}

    def remove_document(self, filename: str) -> bool:
        """删除文档"""
        if filename in self.documents:
            del self.documents[filename]
            if filename in self.document_metadata:
                del self.document_metadata[filename]
            return True
        return False

    def get_stats(self) -> dict:
        """获取服务统计信息"""
        total_words = sum(meta.get('word_count', 0) for meta in self.document_metadata.values())
        total_tables = sum(len(meta.get('tables', [])) for meta in self.document_metadata.values())

        return {
            "total_documents": len(self.documents),
            "total_words": total_words,
            "total_tables": total_tables,
            "documents": list(self.documents.keys())
        }


# 全局单例
_long_context_service_instance: Optional[SimpleLongContextService] = None


def get_long_context_service() -> SimpleLongContextService:
    """获取全局服务实例"""
    global _long_context_service_instance

    # 首先尝试从 sys.modules 获取（由 main.py 注册的实例）
    if '_long_context_service_instance' in sys.modules:
        return sys.modules['_long_context_service_instance']

    # 如果没有，使用本地实例
    if _long_context_service_instance is None:
        _long_context_service_instance = SimpleLongContextService()
    return _long_context_service_instance